"""Input parsers for investor source evidence."""

from __future__ import annotations

import csv
import importlib.util
from html.parser import HTMLParser
import json
import os
import re
import shutil
import subprocess
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence

from .classifier import classify_record, should_keep_event
from .events import build_event, build_gap_event
from .profiles import get_profile


TEXT_EXTENSIONS = {".md", ".markdown", ".txt", ".html", ".htm", ".eml", ".ics"}
TABLE_EXTENSIONS = {".csv", ".tsv"}
JSON_EXTENSIONS = {".json", ".jsonl", ".ndjson"}
CONTENT_EXTRACT_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xlsm", ".xls", ".pptx"}
IMAGE_METADATA_ONLY_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".heic", ".heif"}
METADATA_ONLY_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".pptx"} | IMAGE_METADATA_ONLY_EXTENSIONS
RESEARCH_DOCUMENT_EXTENSIONS = (
    TEXT_EXTENSIONS
    | TABLE_EXTENSIONS
    | JSON_EXTENSIONS
    | CONTENT_EXTRACT_EXTENSIONS
    | METADATA_ONLY_EXTENSIONS
    | {".doc", ".ppt", ".key", ".numbers", ".bmp", ".gif", ".heic", ".heif"}
)
MAX_EXTRACTED_CHARS = 20000


class DocumentExtractError(RuntimeError):
    def __init__(self, code: str, *, parser: str) -> None:
        super().__init__(code)
        self.code = code
        self.parser = parser


class CollectionResult:
    def __init__(self, *, events: List[Dict[str, Any]], audit: Dict[str, Any]) -> None:
        self.events = events
        self.audit = audit


def collect_events(
    source_id: str,
    inputs: Iterable[str],
    *,
    collected_at: Optional[str] = None,
    include_content: bool = False,
    include_image_ocr: bool = False,
    limit: Optional[int] = None,
    min_score: float = 0.30,
    include_non_matches: bool = False,
    allow_chats: Optional[Sequence[str]] = None,
    deny_chats: Optional[Sequence[str]] = None,
    allow_senders: Optional[Sequence[str]] = None,
    deny_senders: Optional[Sequence[str]] = None,
) -> List[Dict[str, Any]]:
    return collect_events_with_audit(
        source_id,
        inputs,
        collected_at=collected_at,
        include_content=include_content,
        include_image_ocr=include_image_ocr,
        limit=limit,
        min_score=min_score,
        include_non_matches=include_non_matches,
        allow_chats=allow_chats,
        deny_chats=deny_chats,
        allow_senders=allow_senders,
        deny_senders=deny_senders,
    ).events


def collect_events_with_audit(
    source_id: str,
    inputs: Iterable[str],
    *,
    collected_at: Optional[str] = None,
    include_content: bool = False,
    include_image_ocr: bool = False,
    limit: Optional[int] = None,
    min_score: float = 0.30,
    include_non_matches: bool = False,
    allow_chats: Optional[Sequence[str]] = None,
    deny_chats: Optional[Sequence[str]] = None,
    allow_senders: Optional[Sequence[str]] = None,
    deny_senders: Optional[Sequence[str]] = None,
) -> CollectionResult:
    get_profile(source_id)
    input_list = list(inputs)
    input_resolution = resolve_input_paths(input_list)
    paths = input_resolution["paths"]
    source_policy = build_source_policy(
        allow_chats=allow_chats,
        deny_chats=deny_chats,
        allow_senders=allow_senders,
        deny_senders=deny_senders,
    )
    audit = initial_collection_audit(
        source_id,
        input_list,
        paths,
        include_content=include_content,
        include_image_ocr=include_image_ocr,
        limit=limit,
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        input_resolution=input_resolution,
    )
    if not paths:
        events = [build_gap_event(source_id, collected_at=collected_at)]
        finalize_collection_audit(audit, events, parsed_count=0)
        return CollectionResult(events=events, audit=audit)

    events: List[Dict[str, Any]] = []
    parsed_count = 0
    for path in paths:
        if limit is not None and len(events) >= limit:
            audit["limit_reached"] = True
            break
        parsed = parse_path(
            source_id,
            path,
            collected_at=collected_at,
            include_content=include_content,
            include_image_ocr=include_image_ocr,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        parsed_count += len(parsed.candidates)
        events_to_add = parsed.events
        if limit is not None:
            remaining = max(limit - len(events), 0)
            if len(events_to_add) > remaining:
                audit["limit_reached"] = True
                events_to_add = events_to_add[:remaining]
                mark_last_path_result_limit(audit, path, emitted_event_count=len(events_to_add))
        events.extend(events_to_add)
        if limit is not None and len(events) >= limit:
            audit["limit_reached"] = True
            finalize_collection_audit(audit, events, parsed_count=parsed_count)
            return CollectionResult(events=events, audit=audit)
    if not events:
        source_policy_filtered = int((audit.get("source_policy") or {}).get("filtered_candidate_count") or 0)
        if parsed_count == 0:
            reason = "no_readable_input"
        elif source_policy_filtered >= parsed_count:
            reason = "source_policy_filtered_all"
        else:
            reason = "no_investment_evidence_matched"
        events = [build_gap_event(source_id, collected_at=collected_at, reason=reason)]
    finalize_collection_audit(audit, events, parsed_count=parsed_count)
    return CollectionResult(events=events, audit=audit)


def resolve_input_paths(inputs: Iterable[str]) -> Dict[str, Any]:
    paths: List[Path] = []
    path_results: List[Dict[str, Any]] = []
    skipped_reason_counts: Counter[str] = Counter()
    skipped_file_count = 0
    input_missing_count = 0
    requested_inputs: List[str] = []
    for raw in inputs:
        path = Path(raw).expanduser()
        requested_inputs.append(str(path))
        if not path.exists():
            input_missing_count += 1
            skipped_reason_counts["input_missing"] += 1
            path_results.append(path_result(path, status="missing", reason="input_missing"))
            continue
        if path.is_dir():
            for child in sorted(path.rglob("*")):
                if not child.is_file():
                    continue
                if child.name.startswith("."):
                    skipped_file_count += 1
                    skipped_reason_counts["hidden_file"] += 1
                    path_results.append(path_result(child, status="skipped", reason="hidden_file"))
                    continue
                paths.append(child)
        elif path.is_file():
            paths.append(path)
        else:
            skipped_reason_counts["unsupported_input_kind"] += 1
            path_results.append(path_result(path, status="skipped", reason="unsupported_input_kind"))
    return {
        "paths": paths,
        "requested_inputs": requested_inputs,
        "input_missing_count": input_missing_count,
        "skipped_file_count": skipped_file_count,
        "skipped_reason_counts": skipped_reason_counts,
        "path_results": path_results,
    }


def iter_input_paths(inputs: Iterable[str]) -> Iterator[Path]:
    yield from resolve_input_paths(inputs)["paths"]


def parse_path(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    include_content: bool,
    include_image_ocr: bool,
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]] = None,
) -> "ParseResult":
    suffix = path.suffix.lower()
    if audit is not None:
        audit_counter(audit, "extension_counts")[suffix or "<none>"] += 1
    if source_id == "research-documents" and suffix not in RESEARCH_DOCUMENT_EXTENSIONS:
        if audit is not None:
            audit["skipped_file_count"] += 1
            audit_counter(audit, "skipped_extension_counts")[suffix or "<none>"] += 1
            audit_counter(audit, "skipped_reason_counts")["unsupported_extension"] += 1
            audit.setdefault("path_results", []).append(path_result(path, status="skipped", reason="unsupported_extension"))
        return ParseResult(candidates=[], events=[])
    if suffix in TABLE_EXTENSIONS:
        parsed = parse_table(
            source_id,
            path,
            collected_at=collected_at,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        record_path_parse_result(audit, path, parsed, parser="csv")
        return parsed
    if suffix in JSON_EXTENSIONS:
        parsed = parse_json_like(
            source_id,
            path,
            collected_at=collected_at,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        record_path_parse_result(audit, path, parsed, parser="json")
        return parsed
    if suffix in TEXT_EXTENSIONS:
        event = parse_text_file(
            source_id,
            path,
            collected_at=collected_at,
            include_content=include_content,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        parsed = ParseResult(candidates=[path], events=[event] if event else [])
        record_path_parse_result(audit, path, parsed, parser="text")
        return parsed
    if include_content and suffix in CONTENT_EXTRACT_EXTENSIONS:
        event = parse_content_file(
            source_id,
            path,
            collected_at=collected_at,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        parsed = ParseResult(candidates=[path], events=[event] if event else [])
        parser = (event.get("raw_ref") or {}).get("parser") if event else "content_extract"
        record_path_parse_result(audit, path, parsed, parser=str(parser or "content_extract"))
        return parsed
    if source_id == "research-documents" and include_image_ocr and suffix in IMAGE_METADATA_ONLY_EXTENSIONS:
        event = parse_image_ocr_file(
            source_id,
            path,
            collected_at=collected_at,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        parsed = ParseResult(candidates=[path], events=[event] if event else [])
        parser = (event.get("raw_ref") or {}).get("parser") if event else "tesseract-ocr"
        record_path_parse_result(audit, path, parsed, parser=str(parser or "tesseract-ocr"))
        return parsed
    event = parse_metadata_file(
        source_id,
        path,
        collected_at=collected_at,
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        audit=audit,
    )
    parsed = ParseResult(candidates=[path], events=[event] if event else [])
    record_path_parse_result(audit, path, parsed, parser="metadata")
    return parsed


class ParseResult:
    def __init__(self, *, candidates: List[Any], events: List[Dict[str, Any]]) -> None:
        self.candidates = candidates
        self.events = events


def initial_collection_audit(
    source_id: str,
    inputs: List[str],
    paths: List[Path],
    *,
    include_content: bool,
    include_image_ocr: bool,
    limit: Optional[int],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    input_resolution: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    input_resolution = input_resolution or {}
    skipped_reason_counts = input_resolution.get("skipped_reason_counts") or Counter()
    return {
        "source_id": source_id,
        "input_count": len(inputs),
        "requested_inputs": input_resolution.get("requested_inputs", [str(Path(raw).expanduser()) for raw in inputs]),
        "resolved_input_file_count": len(paths),
        "input_missing_count": int(input_resolution.get("input_missing_count") or 0),
        "candidate_record_count": 0,
        "matched_event_count": 0,
        "non_matched_event_count": 0,
        "filtered_candidate_count": 0,
        "skipped_file_count": int(input_resolution.get("skipped_file_count") or 0),
        "skipped_reason_counts": skipped_reason_counts,
        "extension_counts": {},
        "skipped_extension_counts": {},
        "parser_counts": {},
        "metadata_only_file_count": 0,
        "screenshot_metadata_only_file_count": 0,
        "ocr_performed": False,
        "include_image_ocr": include_image_ocr,
        "image_ocr_event_count": 0,
        "image_ocr_status_counts": {},
        "content_read_event_count": 0,
        "content_extract_status_counts": {},
        "include_content": include_content,
        "include_non_matches": include_non_matches,
        "min_score": min_score,
        "limit": limit,
        "limit_reached": False,
        "path_results": list(input_resolution.get("path_results") or []),
        "content_extraction_policy": content_extraction_policy(source_id, include_content, include_image_ocr),
        "source_policy": {
            "enabled": source_policy_enabled(source_policy),
            "allow_chats": (source_policy or {}).get("allow_chats", []),
            "deny_chats": (source_policy or {}).get("deny_chats", []),
            "allow_senders": (source_policy or {}).get("allow_senders", []),
            "deny_senders": (source_policy or {}).get("deny_senders", []),
            "filtered_candidate_count": 0,
            "filter_reason_counts": {},
            "policy_does_not_assert_investment_relevance": True,
        },
    }


def finalize_collection_audit(audit: Dict[str, Any], events: List[Dict[str, Any]], *, parsed_count: int) -> None:
    usable_events = [event for event in events if (event.get("data") or {}).get("payload", {}).get("signal_type") != "collector_preflight_gap"]
    audit["candidate_record_count"] = parsed_count
    audit["matched_event_count"] = sum(
        1
        for event in usable_events
        if ((event.get("data") or {}).get("classification") or {}).get("is_investment_evidence")
    )
    audit["non_matched_event_count"] = sum(
        1
        for event in usable_events
        if not ((event.get("data") or {}).get("classification") or {}).get("is_investment_evidence")
    )
    audit["filtered_candidate_count"] = max(0, parsed_count - len(usable_events))

    parser_counts: Counter[str] = Counter()
    content_status_counts: Counter[str] = Counter()
    image_ocr_status_counts: Counter[str] = Counter()
    image_ocr_event_count = 0
    content_read_count = 0
    for event in usable_events:
        raw_ref = event.get("raw_ref") or {}
        payload = (event.get("data") or {}).get("payload") or {}
        parser = raw_ref.get("parser")
        if parser:
            parser_counts[str(parser)] += 1
        if raw_ref.get("content_read"):
            content_read_count += 1
        extract = payload.get("content_extract") if isinstance(payload, dict) else None
        if isinstance(extract, dict):
            status = extract.get("status") or "unknown"
            content_status_counts[str(status)] += 1
            if raw_ref.get("image_ocr_requested"):
                image_ocr_status_counts[str(status)] += 1
        if raw_ref.get("image_ocr_performed"):
            image_ocr_event_count += 1
    audit["parser_counts"] = dict(sorted(parser_counts.items()))
    audit["content_read_event_count"] = content_read_count
    audit["content_extract_status_counts"] = dict(sorted(content_status_counts.items()))
    audit["image_ocr_event_count"] = image_ocr_event_count
    audit["image_ocr_status_counts"] = dict(sorted(image_ocr_status_counts.items()))
    audit["ocr_performed"] = image_ocr_event_count > 0
    audit["extension_counts"] = dict(sorted(audit_counter(audit, "extension_counts").items()))
    audit["skipped_extension_counts"] = dict(sorted(audit_counter(audit, "skipped_extension_counts").items()))
    audit["skipped_reason_counts"] = dict(sorted(audit_counter(audit, "skipped_reason_counts").items()))
    source_policy = audit.get("source_policy") or {}
    if isinstance(source_policy.get("filter_reason_counts"), Counter):
        source_policy["filter_reason_counts"] = dict(sorted(source_policy["filter_reason_counts"].items()))


def content_extraction_policy(source_id: str, include_content: bool, include_image_ocr: bool = False) -> Dict[str, Any]:
    if source_id != "research-documents":
        return {
            "include_content_enabled": include_content,
            "include_image_ocr_enabled": include_image_ocr,
            "applies_to": "generic investor-source input parser",
        }
    ocr_engine = image_ocr_engine()
    return {
        "applies_to": "research-documents lens",
        "input_boundary": "user_selected_files_or_folders_only",
        "generic_filesystem_collector": "metadata_only",
        "content_read_requires_explicit_include_content": True,
        "include_content_enabled": include_content,
        "image_ocr_requires_explicit_include_image_ocr": True,
        "include_image_ocr_enabled": include_image_ocr,
        "image_ocr_engine": ocr_engine if include_image_ocr else None,
        "text_files_read_for_preview_extensions": sorted(TEXT_EXTENSIONS),
        "table_files_read_as_rows_extensions": sorted(TABLE_EXTENSIONS),
        "binary_content_extract_extensions": sorted(CONTENT_EXTRACT_EXTENSIONS),
        "binary_metadata_only_extensions_without_include_content": sorted(METADATA_ONLY_EXTENSIONS),
        "screenshot_metadata_only_extensions": sorted(IMAGE_METADATA_ONLY_EXTENSIONS),
        "screenshots_are_metadata_only_no_ocr": not include_image_ocr,
        "ocr_performed": False,
        "ocr_requires_separate_user_consent_and_adapter": not include_image_ocr,
        "ocr_engine_required": "tesseract",
        "unsupported_extensions_are_skipped": True,
        "preview_char_limit": 1200,
        "extracted_text_char_limit": MAX_EXTRACTED_CHARS,
        "legacy_xls_parser_variants": [
            "openpyxl-renamed-xls",
            "legacy-xls-html",
            "legacy-xls-xml",
            "legacy-xls-delimited",
            "legacy-xls-text",
            "xlrd-biff",
        ],
        "binary_xls_biff_requires_xlrd": True,
        "binary_xls_biff_parser_available": xlrd_available(),
        "binary_xls_without_xlrd_records_extract_failed": True,
        "collector_writes_wiki_directly": False,
    }


def audit_counter(audit: Dict[str, Any], key: str) -> Counter[str]:
    value = audit.get(key)
    if isinstance(value, Counter):
        return value
    counter: Counter[str] = Counter(value or {})
    audit[key] = counter
    return counter


def path_result(path: Path, *, status: str, reason: Optional[str] = None) -> Dict[str, Any]:
    result = {
        "path": str(path),
        "extension": path.suffix.lower() or "<none>",
        "status": status,
    }
    if reason:
        result["reason"] = reason
    return result


def record_path_parse_result(audit: Optional[Dict[str, Any]], path: Path, parsed: ParseResult, *, parser: str) -> None:
    if audit is None:
        return
    suffix = path.suffix.lower()
    result = path_result(path, status="parsed" if parsed.candidates else "no_candidates")
    result.update(
        {
            "parser": parser,
            "candidate_record_count": len(parsed.candidates),
            "emitted_event_count": len(parsed.events),
        }
    )
    if parser == "metadata":
        audit["metadata_only_file_count"] = int(audit.get("metadata_only_file_count") or 0) + 1
    if suffix in IMAGE_METADATA_ONLY_EXTENSIONS:
        ocr_performed = any((event.get("raw_ref") or {}).get("image_ocr_performed") for event in parsed.events)
        ocr_requested = parser == "tesseract-ocr" or any((event.get("raw_ref") or {}).get("image_ocr_requested") for event in parsed.events)
        result["ocr_requested"] = ocr_requested
        result["ocr_performed"] = ocr_performed
        if ocr_performed:
            result["content_policy"] = "image_ocr_explicit_authorized"
        else:
            audit["screenshot_metadata_only_file_count"] = int(audit.get("screenshot_metadata_only_file_count") or 0) + 1
            result["content_policy"] = "image_ocr_requested_but_not_performed" if ocr_requested else "screenshot_metadata_only_no_ocr"
    audit.setdefault("path_results", []).append(result)


def mark_last_path_result_limit(audit: Dict[str, Any], path: Path, *, emitted_event_count: int) -> None:
    for result in reversed(audit.get("path_results") or []):
        if result.get("path") == str(path):
            result["emitted_event_count"] = emitted_event_count
            result["limit_truncated"] = True
            return


def parse_table(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> ParseResult:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    text = path.read_text(encoding="utf-8-sig")
    rows = csv.DictReader(text.splitlines(), delimiter=delimiter)
    events: List[Dict[str, Any]] = []
    candidates: List[Any] = []
    for index, row in enumerate(rows, start=1):
        record = {str(key): value for key, value in row.items() if key is not None}
        candidates.append(record)
        event = candidate_to_event(
            source_id,
            record,
            source_label=str(path),
            raw_ref={"path": str(path), "row": index, "parser": "csv"},
            collected_at=collected_at,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        if event:
            events.append(event)
    return ParseResult(candidates=candidates, events=events)


def parse_json_like(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> ParseResult:
    text = path.read_text(encoding="utf-8-sig").strip()
    if not text:
        return ParseResult(candidates=[], events=[])
    if path.suffix.lower() in {".jsonl", ".ndjson"}:
        items = [json.loads(line) for line in text.splitlines() if line.strip()]
    else:
        loaded = json.loads(text)
        if isinstance(loaded, list):
            items = loaded
        elif isinstance(loaded, dict):
            for key in ("events", "items", "records", "data", "messages"):
                if isinstance(loaded.get(key), list):
                    items = loaded[key]
                    break
            else:
                items = [loaded]
        else:
            items = [{"value": loaded}]

    events: List[Dict[str, Any]] = []
    candidates: List[Any] = []
    for index, item in enumerate(items, start=1):
        record, source_label, raw_ref, event_kind, event_time = normalize_json_candidate(item, path, index)
        candidates.append(record)
        event = candidate_to_event(
            source_id,
            record,
            source_label=source_label,
            raw_ref=raw_ref,
            collected_at=collected_at,
            event_kind=event_kind,
            event_time=event_time,
            min_score=min_score,
            include_non_matches=include_non_matches,
            source_policy=source_policy,
            audit=audit,
        )
        if event:
            events.append(event)
    return ParseResult(candidates=candidates, events=events)


def parse_text_file(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    include_content: bool,
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> Optional[Dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    record: Dict[str, Any] = {
        "title": path.stem,
        "path": str(path),
        "content_preview": text[:1200],
        "byte_size": path.stat().st_size,
    }
    if include_content:
        record["content"] = text
    return candidate_to_event(
        source_id=source_id,
        record=record,
        source_label=str(path),
        raw_ref={"path": str(path), "parser": "text", "byte_size": path.stat().st_size},
        collected_at=collected_at,
        event_kind=kind_for_text_source(source_id),
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        audit=audit,
    )


def parse_content_file(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> Optional[Dict[str, Any]]:
    extracted = extract_document_text(path)
    record: Dict[str, Any] = {
        "title": path.stem,
        "path": str(path),
        "extension": path.suffix.lower(),
        "byte_size": path.stat().st_size,
        "metadata_only": False,
        "content_read": extracted["status"] == "extracted",
        "content_extract": {
            "status": extracted["status"],
            "parser": extracted["parser"],
            "text_length": extracted["text_length"],
            "truncated": extracted["truncated"],
        },
    }
    if extracted["text"]:
        record["content_preview"] = extracted["text"][:1200]
        record["content"] = extracted["text"]
    if extracted["error"]:
        record["content_extract"]["error"] = extracted["error"]
    return candidate_to_event(
        source_id=source_id,
        record=record,
        source_label=str(path),
        raw_ref={
            "path": str(path),
            "parser": extracted["parser"],
            "byte_size": path.stat().st_size,
            "content_read": extracted["status"] == "extracted",
            "content_truncated": extracted["truncated"],
        },
        collected_at=collected_at,
        event_kind="file",
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        audit=audit,
    )


def parse_metadata_file(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> Optional[Dict[str, Any]]:
    record = {
        "title": path.stem,
        "path": str(path),
        "extension": path.suffix.lower(),
        "byte_size": path.stat().st_size,
        "metadata_only": True,
    }
    return candidate_to_event(
        source_id=source_id,
        record=record,
        source_label=str(path),
        raw_ref={"path": str(path), "parser": "metadata", "byte_size": path.stat().st_size},
        collected_at=collected_at,
        event_kind="file",
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        audit=audit,
    )


def parse_image_ocr_file(
    source_id: str,
    path: Path,
    *,
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    source_policy: Optional[Dict[str, Any]],
    audit: Optional[Dict[str, Any]],
) -> Optional[Dict[str, Any]]:
    extracted = extract_image_ocr_text(path)
    content_read = extracted["status"] == "extracted"
    record: Dict[str, Any] = {
        "title": path.stem,
        "path": str(path),
        "extension": path.suffix.lower(),
        "byte_size": path.stat().st_size,
        "metadata_only": not content_read,
        "content_read": content_read,
        "image_ocr_requested": True,
        "image_ocr_performed": content_read,
        "content_extract": {
            "status": extracted["status"],
            "parser": extracted["parser"],
            "text_length": extracted["text_length"],
            "truncated": extracted["truncated"],
            "ocr": True,
        },
    }
    if extracted["text"]:
        record["content_preview"] = extracted["text"][:1200]
        record["content"] = extracted["text"]
    if extracted["error"]:
        record["content_extract"]["error"] = extracted["error"]
    return candidate_to_event(
        source_id=source_id,
        record=record,
        source_label=str(path),
        raw_ref={
            "path": str(path),
            "parser": extracted["parser"],
            "byte_size": path.stat().st_size,
            "content_read": content_read,
            "content_truncated": extracted["truncated"],
            "image_ocr_requested": True,
            "image_ocr_performed": content_read,
        },
        collected_at=collected_at,
        event_kind="file",
        min_score=min_score,
        include_non_matches=include_non_matches,
        source_policy=source_policy,
        audit=audit,
    )


def extract_document_text(path: Path) -> Dict[str, Any]:
    suffix = path.suffix.lower()
    parser = suffix.lstrip(".")
    try:
        if suffix == ".pdf":
            text = extract_pdf_text(path)
            parser = "pdfplumber"
        elif suffix == ".docx":
            text = extract_docx_text(path)
            parser = "python-docx"
        elif suffix in {".xlsx", ".xlsm"}:
            text = extract_xlsx_text(path)
            parser = "openpyxl"
        elif suffix == ".xls":
            text, parser = extract_xls_text_with_parser(path)
        elif suffix == ".pptx":
            text = extract_pptx_text(path)
            parser = "pptx-xml"
        else:
            text = ""
    except DocumentExtractError as exc:
        return {
            "status": "extract_failed",
            "parser": exc.parser,
            "text": "",
            "text_length": 0,
            "truncated": False,
            "error": exc.code,
        }
    except Exception as exc:  # pragma: no cover - dependency/runtime specific
        return {
            "status": "extract_failed",
            "parser": parser,
            "text": "",
            "text_length": 0,
            "truncated": False,
            "error": str(exc) or type(exc).__name__,
        }
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    truncated = len(normalized) > MAX_EXTRACTED_CHARS
    return {
        "status": "extracted" if normalized else "empty",
        "parser": parser,
        "text": normalized[:MAX_EXTRACTED_CHARS],
        "text_length": len(normalized),
        "truncated": truncated,
        "error": None,
    }


def extract_image_ocr_text(path: Path) -> Dict[str, Any]:
    parser = "tesseract-ocr"
    engine = image_ocr_engine()
    if not engine:
        return {
            "status": "ocr_engine_unavailable",
            "parser": parser,
            "text": "",
            "text_length": 0,
            "truncated": False,
            "error": "tesseract_not_found",
        }
    try:
        result = subprocess.run(
            [engine, str(path), "stdout", "-l", os.environ.get("COLLECTORX_TESSERACT_LANG", "chi_sim+eng")],
            text=True,
            capture_output=True,
            timeout=60,
            check=False,
        )
    except Exception as exc:  # pragma: no cover - runtime specific
        return {
            "status": "ocr_failed",
            "parser": parser,
            "text": "",
            "text_length": 0,
            "truncated": False,
            "error": type(exc).__name__,
        }
    if result.returncode != 0:
        return {
            "status": "ocr_failed",
            "parser": parser,
            "text": "",
            "text_length": 0,
            "truncated": False,
            "error": (result.stderr or "").strip()[:200] or f"exit_{result.returncode}",
        }
    normalized = "\n".join(line.strip() for line in result.stdout.splitlines() if line.strip())
    truncated = len(normalized) > MAX_EXTRACTED_CHARS
    return {
        "status": "extracted" if normalized else "empty",
        "parser": parser,
        "text": normalized[:MAX_EXTRACTED_CHARS],
        "text_length": len(normalized),
        "truncated": truncated,
        "error": None,
    }


def image_ocr_engine() -> Optional[str]:
    configured = os.environ.get("COLLECTORX_TESSERACT_CMD")
    if configured:
        expanded = Path(configured).expanduser()
        return str(expanded) if expanded.exists() else None
    return shutil.which("tesseract")


def extract_pdf_text(path: Path) -> str:
    import pdfplumber  # type: ignore

    parts: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:20]:
            parts.append(page.extract_text() or "")
            if sum(len(part) for part in parts) > MAX_EXTRACTED_CHARS:
                break
    return "\n".join(parts)


def extract_docx_text(path: Path) -> str:
    import docx  # type: ignore

    document = docx.Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:20]:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    import openpyxl  # type: ignore

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: List[str] = []
    try:
        for sheet in workbook.worksheets[:10]:
            parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(max_row=200, values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
                if sum(len(part) for part in parts) > MAX_EXTRACTED_CHARS:
                    break
    finally:
        workbook.close()
    return "\n".join(parts)


def extract_xls_text(path: Path) -> str:
    text, _parser = extract_xls_text_with_parser(path)
    return text


def extract_xls_text_with_parser(path: Path) -> tuple[str, str]:
    raw = path.read_bytes()
    if raw.startswith(b"PK"):
        return extract_xlsx_text(path), "openpyxl-renamed-xls"

    text = decode_legacy_text(raw)
    if text:
        lowered = text[:500].lower()
        if "<workbook" in lowered or "urn:schemas-microsoft-com:office:spreadsheet" in lowered:
            return extract_xml_spreadsheet_text(text), "legacy-xls-xml"
        if "<html" in lowered or "<table" in lowered:
            return extract_html_text(text), "legacy-xls-html"
        first_line = text.splitlines()[0] if text.splitlines() else ""
        if "\t" in text or "," in first_line:
            return normalize_tabular_text(text), "legacy-xls-delimited"
        if is_probably_plain_text(text):
            return text.strip(), "legacy-xls-text"

    if not xlrd_available():
        raise DocumentExtractError("xlrd_unavailable_for_binary_xls", parser="xlrd-biff")

    import xlrd  # type: ignore
    workbook = xlrd.open_workbook(str(path), on_demand=True)
    parts: List[str] = []
    try:
        for sheet in workbook.sheets()[:10]:
            parts.append(f"# {sheet.name}")
            for row_index in range(min(sheet.nrows, 200)):
                values = [str(value) for value in sheet.row_values(row_index) if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
                if sum(len(part) for part in parts) > MAX_EXTRACTED_CHARS:
                    break
    finally:
        workbook.release_resources()
    return "\n".join(parts), "xlrd-biff"


def xlrd_available() -> bool:
    if os.environ.get("COLLECTORX_DISABLE_XLRD"):
        return False
    return importlib.util.find_spec("xlrd") is not None


def decode_legacy_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            text = raw.decode(encoding)
        except UnicodeDecodeError:
            continue
        if text.strip():
            return text
    return ""


def is_probably_plain_text(text: str) -> bool:
    sample = text[:4096]
    if not sample.strip():
        return False
    control_chars = sum(1 for char in sample if ord(char) < 32 and char not in "\r\n\t")
    return (control_chars / max(len(sample), 1)) < 0.02


def extract_xml_spreadsheet_text(text: str) -> str:
    root = ET.fromstring(text)
    parts: List[str] = []
    for node in root.iter():
        local_name = node.tag.rsplit("}", 1)[-1].lower()
        if local_name == "data" and node.text and node.text.strip():
            parts.append(node.text.strip())
        if len("\n".join(parts)) > MAX_EXTRACTED_CHARS:
            break
    return "\n".join(parts)


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: List[str] = []

    def handle_data(self, data: str) -> None:
        value = data.strip()
        if value:
            self.parts.append(value)


def extract_html_text(text: str) -> str:
    parser = _HTMLTextParser()
    parser.feed(text)
    return "\n".join(parser.parts)


def normalize_tabular_text(text: str) -> str:
    parts: List[str] = []
    for row in csv.reader(text.splitlines(), delimiter="\t" if "\t" in text else ","):
        values = [value.strip() for value in row if value.strip()]
        if values:
            parts.append(" | ".join(values))
        if sum(len(part) for part in parts) > MAX_EXTRACTED_CHARS:
            break
    return "\n".join(parts)


def extract_pptx_text(path: Path) -> str:
    parts: List[str] = []
    with zipfile.ZipFile(path) as archive:
        slide_names = sorted(
            (name for name in archive.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)),
            key=slide_sort_key,
        )
        for slide_name in slide_names[:80]:
            root = ET.fromstring(archive.read(slide_name))
            for node in root.iter():
                local_name = node.tag.rsplit("}", 1)[-1]
                if local_name == "t" and node.text and node.text.strip():
                    parts.append(node.text.strip())
            if sum(len(part) for part in parts) > MAX_EXTRACTED_CHARS:
                break
    return "\n".join(parts)


def slide_sort_key(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def normalize_json_candidate(item: Any, path: Path, index: int) -> tuple[Dict[str, Any], str, Dict[str, Any], Optional[str], Optional[str]]:
    if not isinstance(item, dict):
        return {"value": item}, str(path), {"path": str(path), "row": index, "parser": "json"}, None, None

    if item.get("schema") == "collectorx.event.v1":
        data = item.get("data") if isinstance(item.get("data"), dict) else {}
        record = {
            **data,
            "upstream_collector": item.get("collector"),
            "upstream_kind": item.get("kind"),
            "upstream_source": item.get("source"),
            "upstream_time": item.get("time"),
        }
        return (
            record,
            str(item.get("source") or path),
            {
                "path": str(path),
                "row": index,
                "parser": "collectorx.event.v1",
                "upstream_event_id": item.get("id"),
                "upstream_raw_ref": item.get("raw_ref") or {},
            },
            item.get("kind"),
            item.get("time"),
        )

    if "data" in item and isinstance(item["data"], dict) and set(item) <= {"id", "source", "data"}:
        record = item["data"]
        source_label = str(item.get("source") or path)
        raw_ref = {"path": str(path), "row": index, "parser": "wechat_collect", "upstream_event_id": item.get("id")}
        return record, source_label, raw_ref, "message", record.get("time")

    return item, str(path), {"path": str(path), "row": index, "parser": "json"}, None, None


def candidate_to_event(
    source_id: str,
    record: Dict[str, Any],
    *,
    source_label: str,
    raw_ref: Dict[str, Any],
    collected_at: Optional[str],
    min_score: float,
    include_non_matches: bool,
    event_kind: Optional[str] = None,
    event_time: Optional[str] = None,
    source_policy: Optional[Dict[str, Any]] = None,
    audit: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    allowed, policy_match = source_policy_match(record, source_label=source_label, source_policy=source_policy)
    if not allowed:
        record_source_policy_filter(audit, policy_match.get("reason", "source_policy_filtered"))
        return None
    classification = classify_record(source_id, record)
    if not should_keep_event(source_id, classification, min_score=min_score, include_non_matches=include_non_matches):
        return None
    event = build_event(
        source_id=source_id,
        source_label=source_label,
        record=record,
        raw_ref=raw_ref,
        collected_at=collected_at,
        event_kind=event_kind,
        event_time=event_time,
        classification=classification,
    )
    if source_policy_enabled(source_policy):
        event["data"]["source_policy"] = policy_match
    return event


def build_source_policy(
    *,
    allow_chats: Optional[Sequence[str]],
    deny_chats: Optional[Sequence[str]],
    allow_senders: Optional[Sequence[str]],
    deny_senders: Optional[Sequence[str]],
) -> Dict[str, Any]:
    return {
        "allow_chats": split_patterns(allow_chats),
        "deny_chats": split_patterns(deny_chats),
        "allow_senders": split_patterns(allow_senders),
        "deny_senders": split_patterns(deny_senders),
    }


def split_patterns(values: Optional[Sequence[str]]) -> List[str]:
    out: List[str] = []
    for value in values or []:
        for part in str(value).split(","):
            part = part.strip()
            if part:
                out.append(part)
    return stable_unique(out)


def stable_unique(values: Iterable[str]) -> List[str]:
    seen = set()
    out = []
    for value in values:
        if value in seen:
            continue
        seen.add(value)
        out.append(value)
    return out


def source_policy_enabled(source_policy: Optional[Dict[str, Any]]) -> bool:
    if not source_policy:
        return False
    return any(source_policy.get(key) for key in ("allow_chats", "deny_chats", "allow_senders", "deny_senders"))


def source_policy_match(record: Dict[str, Any], *, source_label: str, source_policy: Optional[Dict[str, Any]]) -> tuple[bool, Dict[str, Any]]:
    if not source_policy_enabled(source_policy):
        return True, {"enabled": False}

    chat_surface = searchable_field_surface(
        record,
        source_label,
        ("chat", "chat_name", "conversation", "group", "source", "upstream_source", "会话", "群", "联系人"),
    )
    sender_surface = searchable_field_surface(
        record,
        source_label,
        ("sender", "from", "author", "speaker", "发送人", "发件人", "说话人"),
    )

    deny_chat = first_pattern_hit(source_policy.get("deny_chats", []), chat_surface)
    if deny_chat:
        return False, {"enabled": True, "allowed": False, "reason": "deny_chat", "matched_pattern": deny_chat}
    deny_sender = first_pattern_hit(source_policy.get("deny_senders", []), sender_surface)
    if deny_sender:
        return False, {"enabled": True, "allowed": False, "reason": "deny_sender", "matched_pattern": deny_sender}

    allow_chat_patterns = source_policy.get("allow_chats", [])
    allow_sender_patterns = source_policy.get("allow_senders", [])
    chat_hit = first_pattern_hit(allow_chat_patterns, chat_surface)
    sender_hit = first_pattern_hit(allow_sender_patterns, sender_surface)
    if allow_chat_patterns and not chat_hit:
        return False, {"enabled": True, "allowed": False, "reason": "allow_chat_not_matched"}
    if allow_sender_patterns and not sender_hit:
        return False, {"enabled": True, "allowed": False, "reason": "allow_sender_not_matched"}

    return True, {
        "enabled": True,
        "allowed": True,
        "matched_allow_chat": chat_hit,
        "matched_allow_sender": sender_hit,
        "policy_does_not_assert_investment_relevance": True,
    }


def searchable_field_surface(record: Dict[str, Any], source_label: str, keys: Iterable[str]) -> str:
    parts = [source_label]
    for key in keys:
        value = record.get(key)
        if value not in (None, ""):
            parts.append(str(value))
    return "\n".join(parts).lower()


def first_pattern_hit(patterns: Iterable[str], surface: str) -> Optional[str]:
    for pattern in patterns:
        if pattern.lower() in surface:
            return pattern
    return None


def record_source_policy_filter(audit: Optional[Dict[str, Any]], reason: str) -> None:
    if not audit:
        return
    policy = audit.get("source_policy")
    if not isinstance(policy, dict):
        return
    policy["filtered_candidate_count"] = int(policy.get("filtered_candidate_count") or 0) + 1
    counts = policy.get("filter_reason_counts")
    if not isinstance(counts, Counter):
        counts = Counter(counts or {})
        policy["filter_reason_counts"] = counts
    counts[reason] += 1


def kind_for_text_source(source_id: str) -> str:
    default = get_profile(source_id)["default_kind"]
    if default in {"message", "email", "task", "calendar"}:
        return default
    if source_id in {"research-documents", "wechat-article-favorites"}:
        return "file"
    return "note"
