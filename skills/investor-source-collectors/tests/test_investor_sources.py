#!/usr/bin/env python3
"""Tests for investor-source-collectors."""

from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parents[1]
SCRIPT = ROOT / "scripts" / "investor_sources.py"
PACKAGE_VALIDATOR = REPO_ROOT / "tools" / "validate_collector_package.py"


def run_cli(*args: str, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run([sys.executable, str(SCRIPT), *args], text=True, capture_output=True, check=True, env=env)


def run_package_validator(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run([sys.executable, str(PACKAGE_VALIDATOR), *args], text=True, capture_output=True, check=True)


def test_list_sources_contains_all_priorities() -> None:
    result = run_cli("list-sources", "--format", "json")
    profiles = json.loads(result.stdout)
    ids = {profile["id"] for profile in profiles}
    assert "wechat-investment-dialogue" in ids
    assert "xueqiu-investor-activity" in ids
    assert "social-investment-influence" in ids
    wechat_profile = next(profile for profile in profiles if profile["id"] == "wechat-investment-dialogue")
    assert wechat_profile["source_policy"]["supports_allow_chat"] is True
    assert wechat_profile["source_policy"]["policy_does_not_assert_investment_relevance"] is True
    research_profile = next(profile for profile in profiles if profile["id"] == "research-documents")
    assert research_profile["document_scope_policy"]["supports_allow_extension"] is True
    assert research_profile["document_scope_policy"]["supports_allow_keyword"] is True
    assert research_profile["document_scope_policy"]["policy_does_not_assert_investment_relevance"] is True
    email_profile = next(profile for profile in profiles if profile["id"] == "email-research")
    assert email_profile["email_research_scope_policy"]["supports_allow_sender_domain"] is True
    assert email_profile["email_research_scope_policy"]["supports_allow_email_surface"] is True
    assert email_profile["email_research_scope_policy"]["policy_does_not_assert_investment_relevance"] is True
    priorities = {profile["priority"] for profile in profiles}
    assert {"P0", "P1", "P2"}.issubset(priorities)
    classes = {profile["collector_class"] for profile in profiles}
    assert {"vertical", "lens"}.issubset(classes)


def test_collect_xueqiu_csv_outputs_event_and_evidence() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        csv_path = root / "xueqiu.csv"
        out_dir = root / "out"
        csv_path.write_text("code,name,group,note\n600519,贵州茅台,白酒,长期观察\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "xueqiu-investor-activity",
            "--input",
            str(csv_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )
        event_path = out_dir / "lake" / "xueqiu-investor-activity" / "events.jsonl"
        events = [json.loads(line) for line in event_path.read_text(encoding="utf-8").splitlines()]
        assert len(events) == 1
        assert events[0]["schema"] == "collectorx.event.v1"
        assert events[0]["kind"] == "watchlist"
        assert events[0]["data"]["normalized"]["symbol"] == "600519"
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        assert evidence["schema"] == "finclaw.investor_wiki_evidence.v1"
        assert evidence["coverage_summary"]["subdimension_count"] == 20


def test_email_research_scope_policy_does_not_filter_other_lenses() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        csv_path = root / "xueqiu.csv"
        out_dir = root / "out"
        csv_path.write_text("code,name,group,note\n600519,贵州茅台,白酒,长期观察\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "xueqiu-investor-activity",
            "--input",
            str(csv_path),
            "--out-dir",
            str(out_dir),
            "--allow-email-sender-domain",
            "broker.example",
            "--allow-email-keyword",
            "半导体",
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )

        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "xueqiu-investor-activity" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        assert events[0]["data"]["normalized"]["symbol"] == "600519"
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["source_policy"]["enabled"] is False
        assert audit["source_policy"]["email_research_scope_applies"] is False
        assert audit["source_policy"]["allow_email_sender_domains"] == []
        assert audit["email_research_scope_policy"]["enabled"] is False


def test_collect_without_input_writes_gap_event() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        out_dir = Path(tmp) / "out"
        run_cli(
            "collect",
            "--source",
            "china-wealth-assets",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "needs_source_authorization_or_input"
        event = json.loads((out_dir / "lake" / "china-wealth-assets" / "events.jsonl").read_text(encoding="utf-8").splitlines()[0])
        assert event["data"]["payload"]["signal_type"] == "collector_preflight_gap"


def test_wechat_lens_keeps_only_investment_dialogue() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "wechat.json"
        out_dir = root / "out"
        source_path.write_text(
            json.dumps(
                [
                    {
                        "id": "wx-1",
                        "source": "跟投研朋友在 2026-07-07 的微信聊天",
                        "data": {
                            "chat": "投研朋友",
                            "sender": "我",
                            "time": "2026-07-07 10:00:00",
                            "text": "准备买入贵州茅台，仓位先到10%，等财报后再复盘。",
                        },
                    },
                    {
                        "id": "wx-2",
                        "source": "跟朋友在 2026-07-07 的微信聊天",
                        "data": {
                            "chat": "朋友",
                            "sender": "朋友",
                            "time": "2026-07-07 11:00:00",
                            "text": "今晚吃饭吗？",
                        },
                    },
                ],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        run_cli(
            "collect",
            "--source",
            "wechat-investment-dialogue",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )
        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "wechat-investment-dialogue" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        assert events[0]["kind"] == "message"
        assert events[0]["data"]["payload"]["text"].startswith("准备买入")
        assert events[0]["data"]["classification"]["is_investment_evidence"] is True
        assert "matched_trade_action_terms" in events[0]["data"]["classification"]["reasons"]
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        proof = manifest["wechat_dialogue_boundary_proof"]
        assert proof["proof_level"] == "authorized_wechat_dialogue_partial"
        assert proof["event_count"] == 1
        assert proof["candidate_record_count"] == 2
        assert proof["matched_event_count"] == 1
        assert proof["filtered_candidate_count"] == 1
        assert proof["complete_wechat_history_claimed"] is False
        assert proof["complete_dialogue_context_claimed"] is False
        assert proof["raw_wechat_database_access"] is False
        surface = manifest["lens_surface_summary"]
        assert surface["chat_counts"] == {"投研朋友": 1}
        assert surface["sender_counts"] == {"我": 1}
        assert surface["wechat_dialogue_surface_counts"]["trade_intention"] == 1
        assert surface["wechat_dialogue_surface_counts"]["position_sizing"] == 1
        assert surface["wechat_dialogue_surface_counts"]["research_discussion"] == 1
        assert surface["wechat_dialogue_surface_counts"]["review_reflection"] == 1
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["wechat-investment-dialogue"]
        assert evidence_surface["wechat_dialogue_surface_counts"]["trade_intention"] == 1


def test_wechat_lens_source_policy_allows_and_denies_chats_and_senders() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "wechat-events.jsonl"
        out_dir = root / "out"
        records = [
            {
                "schema": "collectorx.event.v1",
                "id": "wechat:allowed",
                "collector": "wechat",
                "source": "微信群「投资讨论群」· 2026-07-08 的聊天",
                "owner_scope": "personal",
                "kind": "message",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "chat": "投资讨论群",
                    "sender": "我",
                    "sender_is_owner": True,
                    "text": "准备买入贵州茅台，先看财报和估值再决定仓位。",
                },
                "raw_ref": {"chat": "投资讨论群"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat:outside-chat",
                "collector": "wechat",
                "source": "跟普通朋友在 2026-07-08 的微信聊天",
                "owner_scope": "personal",
                "kind": "message",
                "time": "2026-07-08T09:30:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "chat": "普通朋友",
                    "sender": "普通朋友",
                    "sender_is_owner": False,
                    "text": "这个股票可以买入吗？估值看起来不贵。",
                },
                "raw_ref": {"chat": "普通朋友"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat:denied-sender",
                "collector": "wechat",
                "source": "微信群「投资讨论群」· 2026-07-08 的聊天",
                "owner_scope": "personal",
                "kind": "message",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "chat": "投资讨论群",
                    "sender": "营销号",
                    "sender_is_owner": False,
                    "text": "立即买入这只股票，马上翻倍。",
                },
                "raw_ref": {"chat": "投资讨论群"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(record, ensure_ascii=False) for record in records) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "wechat-investment-dialogue",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--allow-chat",
            "投资讨论群",
            "--deny-sender",
            "营销号",
            "--collected-at",
            "2026-07-08T12:10:00+08:00",
        )
        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "wechat-investment-dialogue" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        assert events[0]["raw_ref"]["upstream_event_id"] == "wechat:allowed"
        assert events[0]["data"]["source_policy"]["allowed"] is True
        assert events[0]["data"]["source_policy"]["matched_allow_chat"] == "投资讨论群"
        assert events[0]["data"]["source_policy"]["policy_does_not_assert_investment_relevance"] is True

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        source_policy = manifest["collection_audit"]["source_policy"]
        assert source_policy["enabled"] is True
        assert source_policy["allow_chats"] == ["投资讨论群"]
        assert source_policy["deny_senders"] == ["营销号"]
        assert source_policy["filtered_candidate_count"] == 2
        assert source_policy["filter_reason_counts"] == {
            "allow_chat_not_matched": 1,
            "deny_sender": 1,
        }
        proof = manifest["wechat_dialogue_boundary_proof"]
        assert proof["proof_level"] == "authorized_wechat_dialogue_with_source_policy"
        assert proof["source_policy_boundary"]["enabled"] is True
        assert proof["source_policy_boundary"]["allow_chats"] == ["投资讨论群"]
        assert proof["source_policy_boundary"]["deny_senders"] == ["营销号"]
        assert proof["source_policy_boundary"]["filtered_candidate_count"] == 2
        dialogue = proof["dialogue_boundary"]
        assert dialogue["chat_count"] == 1
        assert dialogue["sender_count"] == 1
        assert dialogue["owner_message_count"] == 1
        assert dialogue["group_chat_event_count"] == 1
        assert dialogue["events_with_source_policy"] == 1
        assert dialogue["wechat_dialogue_surface_counts"]["trade_intention"] == 1
        assert dialogue["wechat_dialogue_surface_counts"]["buy_sell_reason"] == 1
        assert dialogue["wechat_dialogue_surface_counts"]["position_sizing"] == 1
        assert dialogue["wechat_dialogue_surface_counts"]["research_discussion"] == 1


def test_wechat_lens_source_policy_filtered_all_has_explicit_gap() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "wechat-events.jsonl"
        out_dir = root / "out"
        record = {
            "schema": "collectorx.event.v1",
            "id": "wechat:outside-chat",
            "collector": "wechat",
            "source": "跟普通朋友在 2026-07-08 的微信聊天",
            "owner_scope": "personal",
            "kind": "message",
            "time": "2026-07-08T09:30:00+08:00",
            "collected_at": "2026-07-08T12:00:00+08:00",
            "data": {
                "chat": "普通朋友",
                "sender": "普通朋友",
                "sender_is_owner": False,
                "text": "这个股票可以买入吗？估值看起来不贵。",
            },
            "raw_ref": {"chat": "普通朋友"},
            "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
        }
        source_path.write_text(json.dumps(record, ensure_ascii=False) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "wechat-investment-dialogue",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--allow-chat",
            "投资讨论群",
            "--collected-at",
            "2026-07-08T12:10:00+08:00",
        )
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "source_policy_filtered_all"
        proof = manifest["wechat_dialogue_boundary_proof"]
        assert proof["proof_level"] == "source_policy_filtered_all"
        assert proof["can_enter_finclaw"] is False
        assert proof["event_count"] == 0
        assert proof["candidate_record_count"] == 1
        assert proof["source_policy_boundary"]["filtered_candidate_count"] == 1
        assert proof["dialogue_boundary"]["event_count"] == 0
        event = json.loads((out_dir / "lake" / "wechat-investment-dialogue" / "events.jsonl").read_text(encoding="utf-8").splitlines()[0])
        assert event["data"]["payload"]["gap"] == "source_policy_filtered_all"
        assert event["wiki_targets"] == []


def test_lens_without_investment_match_does_not_fill_wiki_coverage() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "casual.json"
        out_dir = root / "out"
        source_path.write_text(
            json.dumps(
                [
                    {
                        "id": "wx-1",
                        "source": "跟朋友在 2026-07-07 的微信聊天",
                        "data": {
                            "chat": "朋友",
                            "sender": "朋友",
                            "time": "2026-07-07 11:00:00",
                            "text": "今晚吃饭吗？",
                        },
                    }
                ],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        run_cli(
            "collect",
            "--source",
            "wechat-investment-dialogue",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "no_investment_evidence_matched"
        assert manifest["collection_readiness"]["can_claim_complete_source_collection"] is False
        proof = manifest["wechat_dialogue_boundary_proof"]
        assert proof["proof_level"] == "no_usable_investment_dialogue_after_filter"
        assert proof["event_count"] == 0
        assert proof["candidate_record_count"] == 1
        assert proof["can_enter_finclaw"] is False
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        assert evidence["generated_from"]["event_count"] == 0
        assert evidence["coverage_summary"]["usable_for_wiki_now"] == []


def test_email_research_reads_upstream_collectorx_event() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "email-events.jsonl"
        out_dir = root / "out"
        email_event = {
            "schema": "collectorx.event.v1",
            "id": "email:fixture",
            "collector": "email",
            "source": "IMAP 邮件",
            "owner_scope": "personal",
            "kind": "email",
            "time": "2026-07-07T08:00:00+08:00",
            "collected_at": "2026-07-07T15:00:00+08:00",
            "data": {
                "from": "某证券研究所 <research@example.com>",
                "subject": "晨会纪要：半导体行业跟踪",
                "body_preview": "今日晨会讨论半导体行业景气度和重点公司财报。",
            },
            "raw_ref": {"imap_uid": "1", "folder": "INBOX"},
            "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
        }
        source_path.write_text(json.dumps(email_event, ensure_ascii=False) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "email-research",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-07T15:00:00+08:00",
        )
        event = json.loads((out_dir / "lake" / "email-research" / "events.jsonl").read_text(encoding="utf-8").splitlines()[0])
        assert event["kind"] == "email"
        assert event["raw_ref"]["parser"] == "collectorx.event.v1"
        assert event["raw_ref"]["upstream_event_id"] == "email:fixture"
        assert event["data"]["classification"]["confidence"] >= 0.3
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["source_collection_scope"] == "partial_authorized_input"
        assert manifest["collection_readiness"]["can_claim_complete_source_collection"] is False


def test_email_research_matches_research_attachment_filename() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "email-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "email:attachment-research",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T08:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "from": "contact@example.com",
                    "to": "owner@example.com",
                    "subject": "请查收",
                    "body_preview": "附件供参考。",
                    "attachment_refs": [{"filename": "半导体深度报告.pdf", "content_type": "application/pdf"}],
                    "has_attachments": True,
                },
                "raw_ref": {"path": "mail-export.zip::nested/research.eml"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:attachment-casual",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "from": "friend@example.com",
                    "to": "owner@example.com",
                    "subject": "照片",
                    "body_preview": "周末照片。",
                    "attachment_refs": [{"filename": "holiday.jpg", "content_type": "image/jpeg"}],
                    "has_attachments": True,
                },
                "raw_ref": {"path": "mail-export.zip::nested/photo.eml"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "email-research",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T12:10:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "email-research" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 1
        classification = lens_events[0]["data"]["classification"]
        assert "matched_research_attachment_filename" in classification["reasons"]
        assert lens_events[0]["data"]["payload"]["attachment_refs"][0]["filename"] == "半导体深度报告.pdf"


def test_email_research_reports_surface_and_boundary_proof() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "email-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "email:morning",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T08:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "INBOX",
                    "from": "Broker Research <research@broker.example>",
                    "to": "owner@example.com",
                    "subject": "晨会纪要：半导体行业跟踪",
                    "body_preview": "今日晨会讨论半导体行业财报和估值。",
                    "attachment_refs": [{"filename": "半导体晨会纪要.pdf", "content_type": "application/pdf"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<morning@example.com>", "folder": "INBOX"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:roadshow",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "IR",
                    "from": "IR <ir@company.example>",
                    "to": "owner@example.com",
                    "subject": "路演邀请：业绩说明会",
                    "body_preview": "投资者关系团队邀请参加调研交流会。",
                    "attachment_refs": [{"filename": "roadshow.ics", "content_type": "text/calendar"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<roadshow@example.com>", "folder": "IR"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:alert",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "Alerts",
                    "from": "Alerts <alerts@example.com>",
                    "to": "owner@example.com",
                    "subject": "公告提醒：组合持仓公司财报披露",
                    "body_preview": "年报公告和风险提示。",
                },
                "raw_ref": {"message_id": "<alert@example.com>", "folder": "Alerts"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:casual",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-08T11:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "INBOX",
                    "from": "friend@example.com",
                    "to": "owner@example.com",
                    "subject": "周末照片",
                    "body_preview": "照片见附件。",
                    "attachment_refs": [{"filename": "holiday.jpg", "content_type": "image/jpeg"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<casual@example.com>", "folder": "INBOX"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "email-research",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T12:10:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "email-research" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 3
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "events_collected"
        surface = manifest["lens_surface_summary"]
        assert surface["email_research_surface_counts"]["morning_meeting"] == 1
        assert surface["email_research_surface_counts"]["broker_research_report"] == 1
        assert surface["email_research_surface_counts"]["roadshow_invite"] == 1
        assert surface["email_research_surface_counts"]["company_ir_thread"] == 1
        assert surface["email_research_surface_counts"]["earnings_announcement"] == 3
        assert surface["email_research_surface_counts"]["research_attachment"] == 2
        assert surface["email_research_surface_counts"]["portfolio_alert"] == 1
        assert surface["body_preview_event_count"] == 3
        assert surface["full_body_event_count"] == 0
        assert surface["attachment_ref_event_count"] == 2
        assert surface["attachment_filename_event_count"] == 2
        assert surface["research_attachment_event_count"] == 2
        assert surface["message_id_event_count"] == 3
        assert surface["events_with_time"] == 3
        assert surface["sender_domain_counts"] == {
            "broker.example": 1,
            "company.example": 1,
            "example.com": 1,
        }
        proof = manifest["email_research_boundary_proof"]
        assert proof["proof_level"] == "authorized_email_research_with_research_attachment_refs"
        assert proof["event_count"] == 3
        assert proof["candidate_record_count"] == 4
        assert proof["matched_event_count"] == 3
        assert proof["filtered_candidate_count"] == 1
        assert proof["mailbox_boundary"]["folder_counts"] == {"Alerts": 1, "INBOX": 1, "IR": 1}
        assert proof["content_boundary"]["full_body_in_wiki_by_default"] is False
        assert proof["content_boundary"]["attachment_bodies_collected"] is False
        assert proof["complete_mailbox_claimed"] is False
        assert proof["requires_upstream_email_collector"] is True
        validator = run_package_validator(str(out_dir), "--collector", "email-research", "--require-evidence", "--json")
        assert json.loads(validator.stdout)["valid"] is True


def test_email_research_scope_policy_filters_authorized_research_mail() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "email-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "email:allowed",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-09T08:00:00+08:00",
                "collected_at": "2026-07-09T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "INBOX",
                    "from": "Broker Research <morning@broker.example>",
                    "subject": "晨会纪要：半导体行业跟踪",
                    "body_preview": "今日晨会讨论半导体行业财报和估值。",
                    "attachment_refs": [{"filename": "半导体晨会报告.pdf", "content_type": "application/pdf"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<allowed@example.com>", "folder": "INBOX"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:ir",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-09T09:00:00+08:00",
                "collected_at": "2026-07-09T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "IR",
                    "from": "Company IR <ir@company.example>",
                    "subject": "路演邀请：业绩说明会",
                    "body_preview": "投资者关系团队邀请参加调研交流会。",
                    "attachment_refs": [{"filename": "roadshow.ics", "content_type": "text/calendar"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<ir@example.com>", "folder": "IR"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "email:private",
                "collector": "email",
                "source": "授权邮件导出",
                "owner_scope": "personal",
                "kind": "email",
                "time": "2026-07-09T10:00:00+08:00",
                "collected_at": "2026-07-09T12:00:00+08:00",
                "data": {
                    "mailbox": "owner@example.com",
                    "folder": "INBOX",
                    "from": "Broker Research <strategy@broker.example>",
                    "subject": "私人策略：半导体深度报告",
                    "body_preview": "私人观察，不进入投资分身。",
                    "attachment_refs": [{"filename": "半导体深度报告.pdf", "content_type": "application/pdf"}],
                    "has_attachments": True,
                },
                "raw_ref": {"message_id": "<private@example.com>", "folder": "INBOX"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "email-research",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--allow-email-sender-domain",
            "broker.example",
            "--allow-email-folder",
            "INBOX",
            "--allow-email-surface",
            "broker_research_report",
            "--allow-email-keyword",
            "半导体",
            "--allow-email-attachment",
            "报告",
            "--deny-email-keyword",
            "私人",
            "--collected-at",
            "2026-07-09T12:10:00+08:00",
        )

        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "email-research" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 1
        assert lens_events[0]["raw_ref"]["upstream_event_id"] == "email:allowed"
        assert lens_events[0]["data"]["email_research_scope_policy"]["allowed"] is True
        assert lens_events[0]["data"]["email_research_scope_policy"]["matched_allow_email_sender_domain"] == "broker.example"
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        email_policy = audit["email_research_scope_policy"]
        assert email_policy["enabled"] is True
        assert email_policy["allow_email_sender_domains"] == ["broker.example"]
        assert email_policy["allow_email_surfaces"] == ["broker_research_report"]
        assert email_policy["filtered_candidate_count"] == 2
        assert email_policy["filter_reason_counts"] == {
            "allow_email_sender_domain_not_matched": 1,
            "deny_email_keyword": 1,
        }
        proof = manifest["email_research_boundary_proof"]
        assert proof["authorization_scope_boundary"]["enabled"] is True
        assert proof["authorization_scope_boundary"]["filtered_candidate_count"] == 2
        assert proof["authorization_scope_boundary"]["filtered_all"] is False
        assert proof["candidate_record_count"] == 3
        assert proof["matched_event_count"] == 1
        assert proof["complete_mailbox_claimed"] is False


def test_email_research_scope_policy_filtered_all_gap() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "email-events.jsonl"
        out_dir = root / "out"
        event = {
            "schema": "collectorx.event.v1",
            "id": "email:broker",
            "collector": "email",
            "source": "授权邮件导出",
            "owner_scope": "personal",
            "kind": "email",
            "time": "2026-07-09T08:00:00+08:00",
            "collected_at": "2026-07-09T12:00:00+08:00",
            "data": {
                "mailbox": "owner@example.com",
                "folder": "INBOX",
                "from": "Broker Research <morning@broker.example>",
                "subject": "晨会纪要：半导体行业跟踪",
                "body_preview": "今日晨会讨论半导体行业财报和估值。",
            },
            "raw_ref": {"message_id": "<broker@example.com>", "folder": "INBOX"},
            "privacy": {"sensitive": True, "local_only": True, "contains": ["email"]},
        }
        source_path.write_text(json.dumps(event, ensure_ascii=False) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "email-research",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--allow-email-sender-domain",
            "other.example",
            "--collected-at",
            "2026-07-09T12:10:00+08:00",
        )

        lake_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "email-research" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lake_events) == 1
        assert lake_events[0]["data"]["payload"]["gap"] == "email_research_scope_policy_filtered_all"
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "scope_policy_filtered_all"
        assert manifest["collection_readiness"]["can_enter_finclaw"] is False
        audit = manifest["collection_audit"]
        assert audit["email_research_scope_policy_filtered_all"] is True
        assert audit["email_research_scope_policy"]["filter_reason_counts"] == {"allow_email_sender_domain_not_matched": 1}
        proof = manifest["email_research_boundary_proof"]
        assert proof["proof_level"] == "email_research_scope_policy_filtered_all"
        assert proof["authorization_scope_boundary"]["filtered_all"] is True
        assert proof["can_enter_finclaw"] is False


def test_research_documents_extracts_office_and_pdf_content_when_authorized() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document
        from openpyxl import Workbook
        from reportlab.pdfgen import canvas

        root = Path(tmp)
        out_dir = root / "out"

        xlsx_path = root / "semiconductor-model.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "估值表"
        sheet.append(["公司", "财报", "估值", "风险点"])
        sheet.append(["半导体公司A", "现金流改善", "DCF低估", "库存周期"])
        workbook.save(xlsx_path)

        docx_path = root / "roadshow-notes.docx"
        document = Document()
        document.add_heading("路演纪要", level=1)
        document.add_paragraph("讨论买入理由、财报、估值和安全边际。")
        document.save(docx_path)

        pdf_path = root / "factor-report.pdf"
        pdf = canvas.Canvas(str(pdf_path))
        pdf.drawString(72, 720, "DCF ROE PE PB research report risk review")
        pdf.save()

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--include-content",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:00:00+08:00",
        )
        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 3
        assert {event["data"]["payload"]["extension"] for event in events} == {".docx", ".pdf", ".xlsx"}
        assert all(event["kind"] == "file" for event in events)
        assert all(event["raw_ref"]["content_read"] is True for event in events)
        assert all(event["data"]["payload"]["content_extract"]["status"] == "extracted" for event in events)
        assert any("DCF低估" in event["data"]["payload"].get("content", "") for event in events)
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "events_collected"
        audit = manifest["collection_audit"]
        assert audit["content_extraction_policy"]["include_content_enabled"] is True
        assert audit["content_extraction_policy"]["content_read_requires_explicit_include_content"] is True
        assert audit["content_read_event_count"] == 3
        assert audit["content_extract_status_counts"] == {"extracted": 3}
        assert audit["parser_counts"] == {"openpyxl": 1, "pdfplumber": 1, "python-docx": 1}
        assert audit["filtered_candidate_count"] == 0
        assert audit["requested_inputs"] == [str(root)]
        assert audit["input_missing_count"] == 0
        assert audit["limit_reached"] is False
        assert len(audit["path_results"]) == 3
        assert {item["parser"] for item in audit["path_results"]} == {"openpyxl", "pdfplumber", "python-docx"}
        assert all(item["candidate_record_count"] == 1 for item in audit["path_results"])
        assert all(item["emitted_event_count"] == 1 for item in audit["path_results"])
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "authorized_research_corpus_with_content"
        assert proof["candidate_record_count"] == 3
        assert proof["matched_event_count"] == 3
        assert proof["complete_research_corpus_claimed"] is False
        assert proof["whole_disk_scan_claimed"] is False
        assert proof["public_report_database_crawl_claimed"] is False
        assert proof["content_boundary"]["content_read_event_count"] == 3
        assert proof["content_boundary"]["metadata_only_file_count"] == 0
        assert proof["format_boundary"]["parser_counts"] == {"openpyxl": 1, "pdfplumber": 1, "python-docx": 1}
        surface = manifest["lens_surface_summary"]
        assert surface["extension_counts"] == {".docx": 1, ".pdf": 1, ".xlsx": 1}
        assert surface["content_read_event_count"] == 3
        assert surface["research_document_surface_counts"]["valuation_model"] == 3
        assert surface["research_document_surface_counts"]["research_report"] == 2
        assert surface["research_document_surface_counts"]["table_model"] == 1
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["research-documents"]
        assert evidence_surface["content_read_event_count"] == 3
        assert evidence_surface["research_document_surface_counts"]["valuation_model"] == 3


def test_research_documents_extracts_legacy_xls_and_pptx_when_authorized() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        out_dir = root / "out"

        xls_path = root / "legacy-valuation.xls"
        xls_path.write_text(
            """<?xml version="1.0"?>
<Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet">
  <Worksheet ss:Name="估值表" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet">
    <Table>
      <Row>
        <Cell><Data ss:Type="String">公司</Data></Cell>
        <Cell><Data ss:Type="String">半导体旧模型</Data></Cell>
      </Row>
      <Row>
        <Cell><Data ss:Type="String">DCF 估值</Data></Cell>
        <Cell><Data ss:Type="String">买入理由 安全边际 风险提示</Data></Cell>
      </Row>
    </Table>
  </Worksheet>
</Workbook>
""",
            encoding="utf-8",
        )
        text_xls_path = root / "plain-text-export.xls"
        text_xls_path.write_text(
            "股票研究\n半导体投资计划\n估值 DCF 买入理由 安全边际 风险提示\n",
            encoding="utf-8",
        )

        pptx_path = root / "roadshow-deck.pptx"
        with zipfile.ZipFile(pptx_path, "w") as archive:
            archive.writestr(
                "ppt/slides/slide1.xml",
                """<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody>
    <a:p><a:r><a:t>路演纪要 财报 估值 DCF 买入理由 风险提示</a:t></a:r></a:p>
  </p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>""",
            )

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--include-content",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:10:00+08:00",
        )

        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 3
        assert {event["data"]["payload"]["extension"] for event in events} == {".xls", ".pptx"}
        assert all(event["time"] == "2026-07-08T06:10:00+08:00" for event in events)
        assert all(event["raw_ref"]["content_read"] is True for event in events)
        assert all(event["data"]["payload"]["content_extract"]["status"] == "extracted" for event in events)
        assert any("半导体旧模型" in event["data"]["payload"]["content"] for event in events)
        assert any("半导体投资计划" in event["data"]["payload"]["content"] for event in events)
        assert any("路演纪要" in event["data"]["payload"]["content"] for event in events)

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["content_read_event_count"] == 3
        assert audit["content_extract_status_counts"] == {"extracted": 3}
        assert audit["parser_counts"] == {"legacy-xls-text": 1, "legacy-xls-xml": 1, "pptx-xml": 1}
        assert {item["parser"] for item in audit["path_results"]} == {"legacy-xls-text", "legacy-xls-xml", "pptx-xml"}
        assert ".xls" in audit["content_extraction_policy"]["binary_content_extract_extensions"]
        assert ".pptx" in audit["content_extraction_policy"]["binary_content_extract_extensions"]
        assert audit["content_extraction_policy"]["binary_xls_biff_requires_xlrd"] is True
        assert audit["content_extraction_policy"]["binary_xls_without_xlrd_records_extract_failed"] is True
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "authorized_research_corpus_with_content"
        assert proof["format_boundary"]["parser_counts"] == {"legacy-xls-text": 1, "legacy-xls-xml": 1, "pptx-xml": 1}
        surface = manifest["lens_surface_summary"]
        assert surface["extension_counts"] == {".pptx": 1, ".xls": 2}
        assert surface["research_document_surface_counts"]["valuation_model"] >= 2
        assert surface["research_document_surface_counts"]["research_report"] >= 1
        assert surface["research_document_surface_counts"]["table_model"] >= 1
        validator = run_package_validator(str(out_dir), "--collector", "research-documents", "--require-evidence", "--json")
        assert json.loads(validator.stdout)["valid"] is True


def test_research_documents_binary_xls_without_xlrd_records_explicit_failure() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        out_dir = root / "out"
        binary_xls_path = root / "binary-估值模型.xls"
        binary_xls_path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + (b"\x00" * 2048))
        env = {
            **os.environ,
            "COLLECTORX_DISABLE_XLRD": "1",
        }

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(binary_xls_path),
            "--include-content",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-09T09:00:00+08:00",
            env=env,
        )

        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        event = events[0]
        assert event["kind"] == "file"
        assert event["raw_ref"]["parser"] == "xlrd-biff"
        assert event["raw_ref"]["content_read"] is False
        payload = event["data"]["payload"]
        assert payload["extension"] == ".xls"
        assert payload["content_read"] is False
        assert payload["content_extract"]["status"] == "extract_failed"
        assert payload["content_extract"]["parser"] == "xlrd-biff"
        assert payload["content_extract"]["error"] == "xlrd_unavailable_for_binary_xls"
        assert "content" not in payload

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["content_read_event_count"] == 0
        assert audit["content_extract_status_counts"] == {"extract_failed": 1}
        assert audit["parser_counts"] == {"xlrd-biff": 1}
        assert audit["content_extraction_policy"]["binary_xls_biff_parser_available"] is False
        assert audit["content_extraction_policy"]["binary_xls_without_xlrd_records_extract_failed"] is True
        assert audit["path_results"][0]["parser"] == "xlrd-biff"
        assert audit["path_results"][0]["candidate_record_count"] == 1
        assert audit["path_results"][0]["emitted_event_count"] == 1
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["complete_research_corpus_claimed"] is False
        assert proof["format_boundary"]["parser_counts"] == {"xlrd-biff": 1}
        validator = run_package_validator(str(out_dir), "--collector", "research-documents", "--require-evidence", "--json")
        assert json.loads(validator.stdout)["valid"] is True


def test_research_documents_without_include_content_keeps_binary_metadata_only() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document

        root = Path(tmp)
        out_dir = root / "out"
        docx_path = root / "财报复盘.docx"
        document = Document()
        document.add_paragraph("这段正文包含 DCF 低估和买入理由，但本轮没有授权读取正文。")
        document.save(docx_path)

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:20:00+08:00",
        )
        event = json.loads((out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()[0])
        assert event["kind"] == "file"
        assert event["raw_ref"]["parser"] == "metadata"
        assert event["data"]["payload"]["metadata_only"] is True
        assert "content" not in event["data"]["payload"]
        assert "content_extract" not in event["data"]["payload"]
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["content_extraction_policy"]["include_content_enabled"] is False
        assert audit["content_read_event_count"] == 0
        assert audit["metadata_only_file_count"] == 1
        assert audit["path_results"][0]["parser"] == "metadata"
        assert audit["path_results"][0]["emitted_event_count"] == 1
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "authorized_research_corpus_metadata_only"
        assert proof["content_boundary"]["include_content_enabled"] is False
        assert proof["content_boundary"]["content_read_event_count"] == 0
        assert proof["content_boundary"]["metadata_only_file_count"] == 1
        assert manifest["lens_surface_summary"]["metadata_only_event_count"] == 1


def test_research_documents_filters_broad_titles_and_skips_unsupported_files() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        out_dir = root / "out"
        missing = root / "missing-research-folder"
        (root / "股票计划.png").write_bytes(b"not-an-image-but-metadata-only")
        (root / "研报工具.py").write_text("print('internal helper')\n", encoding="utf-8")

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--input",
            str(missing),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:30:00+08:00",
        )
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "no_investment_evidence_matched"
        audit = manifest["collection_audit"]
        assert audit["candidate_record_count"] == 1
        assert audit["filtered_candidate_count"] == 1
        assert audit["skipped_file_count"] == 1
        assert audit["skipped_extension_counts"] == {".py": 1}
        assert audit["input_count"] == 2
        assert audit["input_missing_count"] == 1
        assert audit["skipped_reason_counts"] == {"input_missing": 1, "unsupported_extension": 1}
        assert audit["screenshot_metadata_only_file_count"] == 1
        assert audit["ocr_performed"] is False
        assert audit["content_extraction_policy"]["screenshots_are_metadata_only_no_ocr"] is True
        assert audit["content_extraction_policy"]["ocr_performed"] is False
        assert audit["content_extraction_policy"]["ocr_requires_separate_user_consent_and_adapter"] is True
        assert audit["content_extraction_policy"]["unsupported_extensions_are_skipped"] is True
        statuses = {(item["extension"], item["status"], item.get("reason")) for item in audit["path_results"]}
        assert (".png", "parsed", None) in statuses
        assert (".py", "skipped", "unsupported_extension") in statuses
        assert ("<none>", "missing", "input_missing") in statuses
        png_result = next(item for item in audit["path_results"] if item["extension"] == ".png")
        assert png_result["content_policy"] == "screenshot_metadata_only_no_ocr"
        assert png_result["ocr_performed"] is False
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "no_usable_research_evidence_after_filter"
        assert proof["can_enter_finclaw"] is False
        assert proof["candidate_record_count"] == 1
        assert proof["filtered_candidate_count"] == 1
        assert proof["content_boundary"]["screenshot_metadata_only_file_count"] == 1
        assert manifest["lens_surface_summary"]["event_count"] == 0
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        assert evidence["coverage_summary"]["usable_for_wiki_now"] == []


def test_research_documents_scope_policy_filters_authorized_surface() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        authorized = root / "authorized"
        blocked = root / "blocked-path"
        authorized.mkdir()
        blocked.mkdir()
        out_dir = root / "out"

        (authorized / "keep-valuation.md").write_text("半导体 DCF 估值 买入理由 风险提示", encoding="utf-8")
        (authorized / "keep-denied.pdf").write_bytes(b"%PDF-1.4 fake")
        (authorized / "keep-extension.docx").write_bytes(b"fake docx")
        (authorized / "keep-table.csv").write_text(
            "title,content\n半导体估值表,半导体 DCF 估值 买入理由 风险提示\n",
            encoding="utf-8",
        )
        (blocked / "keep-path.md").write_text("半导体 DCF 估值 买入理由 风险提示", encoding="utf-8")
        (authorized / "outside-name.md").write_text("半导体 DCF 估值 买入理由 风险提示", encoding="utf-8")
        (authorized / "keep-screen.png").write_bytes(b"fake screenshot")
        (authorized / "keep-keyword.md").write_text("半导体 DCF 估值 禁止采集", encoding="utf-8")
        (authorized / "keep-other.md").write_text("白酒 DCF 估值 买入理由 风险提示", encoding="utf-8")

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--out-dir",
            str(out_dir),
            "--allow-extension",
            "md,csv,png",
            "--deny-extension",
            "pdf",
            "--allow-parser",
            "text,metadata",
            "--allow-path",
            "authorized",
            "--deny-path",
            "blocked-path",
            "--allow-file-name",
            "keep",
            "--allow-research-surface",
            "valuation_model",
            "--deny-research-surface",
            "screenshot_or_image",
            "--allow-keyword",
            "半导体",
            "--deny-keyword",
            "禁止采集",
            "--collected-at",
            "2026-07-09T10:00:00+08:00",
        )

        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        assert events[0]["data"]["payload"]["title"] == "keep-valuation"
        assert events[0]["data"]["document_scope_policy"]["allowed"] is True
        assert events[0]["data"]["document_scope_policy"]["matched_allow_research_surface"] == "valuation_model"

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "events_collected"
        audit = manifest["collection_audit"]
        assert audit["candidate_record_count"] == 9
        assert audit["matched_event_count"] == 1
        assert audit["filtered_candidate_count"] == 8
        policy = audit["document_scope_policy"]
        assert policy["enabled"] is True
        assert policy["allow_extensions"] == [".md", ".csv", ".png"]
        assert policy["deny_extensions"] == [".pdf"]
        assert policy["allow_parsers"] == ["text", "metadata"]
        assert policy["allow_research_surfaces"] == ["valuation_model"]
        assert policy["deny_research_surfaces"] == ["screenshot_or_image"]
        assert policy["filtered_candidate_count"] == 8
        assert policy["filtered_all"] is False
        assert policy["filter_reason_counts"] == {
            "extension_denied": 1,
            "extension_not_allowed": 1,
            "file_name_not_allowed": 1,
            "keyword_denied": 1,
            "keyword_not_allowed": 1,
            "parser_not_allowed": 1,
            "path_denied": 1,
            "research_surface_denied": 1,
        }
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "authorized_research_corpus_with_scope_policy"
        assert proof["authorization_scope_boundary"]["enabled"] is True
        assert proof["authorization_scope_boundary"]["filtered_candidate_count"] == 8
        assert proof["authorization_scope_boundary"]["filtered_all"] is False
        assert proof["can_enter_finclaw"] is True


def test_research_documents_scope_policy_filtered_all_has_explicit_gap() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "keep-valuation.md"
        out_dir = root / "out"
        source_path.write_text("半导体 DCF 估值 买入理由 风险提示", encoding="utf-8")

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--allow-keyword",
            "新能源",
            "--collected-at",
            "2026-07-09T10:10:00+08:00",
        )

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        assert manifest["collection_readiness"]["status"] == "source_policy_filtered_all"
        assert manifest["collection_readiness"]["can_enter_finclaw"] is False
        audit = manifest["collection_audit"]
        assert audit["candidate_record_count"] == 1
        assert audit["document_scope_policy_filtered_all"] is True
        policy = audit["document_scope_policy"]
        assert policy["enabled"] is True
        assert policy["allow_keywords"] == ["新能源"]
        assert policy["filtered_candidate_count"] == 1
        assert policy["filtered_all"] is True
        assert policy["filter_reason_counts"] == {"keyword_not_allowed": 1}
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "source_policy_filtered_all"
        assert proof["event_count"] == 0
        assert proof["authorization_scope_boundary"]["filtered_all"] is True
        assert proof["can_enter_finclaw"] is False
        event = json.loads((out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()[0])
        assert event["data"]["payload"]["gap"] == "source_policy_filtered_all"
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        assert evidence["generated_from"]["event_count"] == 0
        assert evidence["coverage_summary"]["usable_for_wiki_now"] == []


def test_research_documents_image_ocr_requires_explicit_adapter_authorization() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        out_dir = root / "out"
        (root / "估值截图.png").write_bytes(b"fake image bytes")

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:35:00+08:00",
        )

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["content_extraction_policy"]["include_image_ocr_enabled"] is False
        assert audit["content_extraction_policy"]["screenshots_are_metadata_only_no_ocr"] is True
        assert audit["screenshot_metadata_only_file_count"] == 1
        assert audit["ocr_performed"] is False
        png_result = next(item for item in audit["path_results"] if item["extension"] == ".png")
        assert png_result["ocr_requested"] is False
        assert png_result["ocr_performed"] is False
        assert png_result["content_policy"] == "screenshot_metadata_only_no_ocr"


def test_research_documents_image_ocr_extracts_when_explicitly_authorized() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        bin_dir = root / "bin"
        bin_dir.mkdir()
        fake_tesseract = bin_dir / "tesseract"
        fake_tesseract.write_text(
            "#!/bin/sh\n"
            "printf '%s\\n' '半导体 研报 财报 DCF 估值 安全边际 买入理由 风险提示'\n",
            encoding="utf-8",
        )
        fake_tesseract.chmod(0o755)

        image_path = root / "screen.png"
        image_path.write_bytes(b"fake image bytes")
        out_dir = root / "out"
        env = {
            **os.environ,
            "COLLECTORX_TESSERACT_CMD": str(fake_tesseract),
        }

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(image_path),
            "--include-image-ocr",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:36:00+08:00",
            env=env,
        )

        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(events) == 1
        event = events[0]
        assert event["kind"] == "file"
        assert event["raw_ref"]["parser"] == "tesseract-ocr"
        assert event["raw_ref"]["content_read"] is True
        assert event["raw_ref"]["image_ocr_requested"] is True
        assert event["raw_ref"]["image_ocr_performed"] is True
        assert event["data"]["payload"]["metadata_only"] is False
        assert event["data"]["payload"]["content_extract"]["status"] == "extracted"
        assert "DCF" in event["data"]["payload"]["content"]
        assert event["data"]["classification"]["is_investment_evidence"] is True

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert audit["content_extraction_policy"]["include_image_ocr_enabled"] is True
        assert audit["content_extraction_policy"]["image_ocr_engine"] == str(fake_tesseract)
        assert audit["content_extraction_policy"]["screenshots_are_metadata_only_no_ocr"] is False
        assert audit["ocr_performed"] is True
        assert audit["image_ocr_event_count"] == 1
        assert audit["image_ocr_status_counts"] == {"extracted": 1}
        assert audit["content_read_event_count"] == 1
        assert audit["parser_counts"] == {"tesseract-ocr": 1}
        png_result = audit["path_results"][0]
        assert png_result["content_policy"] == "image_ocr_explicit_authorized"
        assert png_result["ocr_requested"] is True
        assert png_result["ocr_performed"] is True
        proof = manifest["research_corpus_boundary_proof"]
        assert proof["proof_level"] == "authorized_research_corpus_with_image_ocr"
        assert proof["content_boundary"]["include_image_ocr_enabled"] is True
        assert proof["content_boundary"]["image_ocr_event_count"] == 1
        assert proof["content_boundary"]["ocr_performed"] is True
        surface = manifest["lens_surface_summary"]
        assert surface["image_ocr_event_count"] == 1
        assert surface["screenshot_or_image_event_count"] == 1
        assert surface["research_document_surface_counts"]["screenshot_or_image"] == 1


def test_research_documents_limit_records_path_truncation() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        out_dir = root / "out"
        (root / "research-records.csv").write_text(
            "title,content\n"
            "半导体深度报告,半导体 DCF 估值 买入理由 风险提示\n"
            "新能源财报复盘,新能源 财报 现金流 估值 安全边际\n",
            encoding="utf-8",
        )

        run_cli(
            "collect",
            "--source",
            "research-documents",
            "--input",
            str(root),
            "--limit",
            "1",
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T06:40:00+08:00",
        )
        events = [
            json.loads(line)
            for line in (out_dir / "lake" / "research-documents" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        audit = manifest["collection_audit"]
        assert len(events) == 1
        assert audit["limit"] == 1
        assert audit["limit_reached"] is True
        assert audit["candidate_record_count"] == 2
        assert audit["matched_event_count"] == 1
        assert audit["path_results"][-1]["limit_truncated"] is True
        assert audit["path_results"][-1]["emitted_event_count"] == 1


def test_task_calendar_lens_keeps_investment_task_and_calendar_only() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "ticktick-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "ticktick:1",
                "collector": "ticktick",
                "source": "滴答清单用户授权任务数据",
                "owner_scope": "personal",
                "kind": "task",
                "time": "2026-07-09T10:00:00+08:00",
                "collected_at": "2026-07-08T01:10:00+08:00",
                "data": {
                    "title": "复盘贵州茅台财报",
                    "content_preview": "看现金流、估值和买入纪律",
                    "project_name": "投资研究",
                },
                "raw_ref": {"task_id": "1"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["task"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "ticktick:2",
                "collector": "ticktick",
                "source": "滴答清单用户授权任务数据",
                "owner_scope": "personal",
                "kind": "task",
                "time": "2026-07-09T12:00:00+08:00",
                "collected_at": "2026-07-08T01:10:00+08:00",
                "data": {"title": "买牛奶", "project_name": "生活"},
                "raw_ref": {"task_id": "2"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["task"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "calendar:1",
                "collector": "calendar",
                "source": "用户授权日历事件",
                "owner_scope": "personal",
                "kind": "calendar",
                "time": "2026-07-10T09:30:00+08:00",
                "collected_at": "2026-07-08T01:10:00+08:00",
                "data": {
                    "title": "贵州茅台财报电话会",
                    "description_preview": "关注现金流、估值、卖出纪律和仓位计划",
                    "calendar_name": "投资日历",
                },
                "raw_ref": {"event_id": "calendar-1"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["calendar"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "calendar:2",
                "collector": "calendar",
                "source": "用户授权日历事件",
                "owner_scope": "personal",
                "kind": "calendar",
                "time": "2026-07-10T12:00:00+08:00",
                "collected_at": "2026-07-08T01:10:00+08:00",
                "data": {"title": "牙医预约", "calendar_name": "生活"},
                "raw_ref": {"event_id": "calendar-2"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["calendar"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "task-calendar-investor",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T01:10:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "task-calendar-investor" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 2
        assert {event["data"]["payload"]["title"] for event in lens_events} == {"复盘贵州茅台财报", "贵州茅台财报电话会"}


def test_task_calendar_lens_reports_planning_surface_from_upstream_events() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "task-calendar-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "ticktick:research",
                "collector": "ticktick",
                "source": "滴答清单用户授权任务数据",
                "owner_scope": "personal",
                "kind": "task",
                "time": "2026-07-09T09:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "ticktick",
                    "title": "研究任务：半导体财报模型",
                    "content_preview": "跟踪财报、估值模型和基本面变化。",
                    "project_name": "投资研究",
                    "due": "2026-07-09T09:00:00+08:00",
                    "time_zone": "Asia/Shanghai",
                    "recurrence": "RRULE:FREQ=WEEKLY;BYDAY=TH",
                    "recurrence_frequency": "weekly",
                    "reminders": ["2026-07-09T08:30:00+08:00"],
                    "has_checklist": True,
                    "checklist_total": 3,
                    "checklist_completed": 2,
                    "checklist_pending": 1,
                    "checklist_completion_rate": 0.6667,
                    "checklist_items": [
                        {"title": "更新财报模型", "is_completed": True},
                        {"title": "复核估值假设", "is_completed": True},
                        {"title": "写交易结论", "is_completed": False},
                    ],
                    "is_completed": False,
                    "is_overdue": False,
                },
                "raw_ref": {"source_app": "ticktick", "task_id": "research"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["task"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "ticktick:trade-plan",
                "collector": "ticktick",
                "source": "滴答清单用户授权任务数据",
                "owner_scope": "personal",
                "kind": "task",
                "time": "2026-07-09T10:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "ticktick",
                    "title": "交易计划：明早买入ETF",
                    "content_preview": "仓位 10%，设置止损和风控检查。",
                    "project_name": "交易计划",
                    "due": "2026-07-09T10:00:00+08:00",
                    "is_completed": False,
                    "is_overdue": False,
                },
                "raw_ref": {"source_app": "ticktick", "task_id": "trade-plan"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["task"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "ticktick:review",
                "collector": "ticktick",
                "source": "滴答清单用户授权任务数据",
                "owner_scope": "personal",
                "kind": "task",
                "time": "2026-07-11T18:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "ticktick",
                    "title": "周五复盘组合回撤",
                    "content_preview": "总结错因、归因和交易纪律。",
                    "project_name": "复盘",
                    "due": "2026-07-11T18:00:00+08:00",
                    "is_completed": False,
                    "is_overdue": False,
                },
                "raw_ref": {"source_app": "ticktick", "task_id": "review"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["task"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "calendar:earnings",
                "collector": "calendar",
                "source": "用户授权日历事件",
                "owner_scope": "personal",
                "kind": "calendar",
                "time": "2026-07-10T09:30:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_platform": "feishu_calendar",
                    "title": "贵州茅台财报电话会",
                    "description_preview": "业绩说明会、调研会议，关注现金流和风险。",
                    "calendar_name": "投资日历",
                    "start": "2026-07-10T09:30:00+08:00",
                    "end": "2026-07-10T10:30:00+08:00",
                    "duration_minutes": 60,
                    "is_multi_day": False,
                    "time_order_valid": True,
                    "meeting_url": "https://example.com/meeting",
                },
                "raw_ref": {"source_platform": "feishu_calendar", "event_id": "earnings"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["calendar"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "calendar:life",
                "collector": "calendar",
                "source": "用户授权日历事件",
                "owner_scope": "personal",
                "kind": "calendar",
                "time": "2026-07-10T12:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {"source_platform": "apple_calendar", "title": "牙医预约", "calendar_name": "生活"},
                "raw_ref": {"source_platform": "apple_calendar", "event_id": "life"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["calendar"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "task-calendar-investor",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T12:30:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "task-calendar-investor" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 4
        assert {event["raw_ref"]["upstream_event_id"] for event in lens_events} == {
            "ticktick:research",
            "ticktick:trade-plan",
            "ticktick:review",
            "calendar:earnings",
        }
        classifications = [event["data"]["classification"] for event in lens_events]
        assert any("research_task" in item["task_calendar_surfaces"] for item in classifications)
        assert any("trade_plan" in item["task_calendar_surfaces"] for item in classifications)
        assert any("review_reminder" in item["task_calendar_surfaces"] for item in classifications)
        assert any("earnings_calendar" in item["task_calendar_surfaces"] for item in classifications)
        assert any("research_meeting" in item["task_calendar_surfaces"] for item in classifications)
        assert any("risk_check" in item["task_calendar_surfaces"] for item in classifications)

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        surface = manifest["lens_surface_summary"]
        assert surface["event_count"] == 4
        assert surface["missing_expected_task_calendar_surfaces"] == []
        assert surface["upstream_collector_counts"] == {"calendar": 1, "ticktick": 3}
        assert surface["kind_counts"] == {"calendar": 1, "task": 3}
        assert surface["source_platform_counts"] == {"feishu_calendar": 1, "ticktick": 3}
        assert surface["task_calendar_surface_counts"]["research_task"] == 2
        assert surface["task_calendar_surface_counts"]["trade_plan"] == 1
        assert surface["task_calendar_surface_counts"]["review_reminder"] == 1
        assert surface["task_calendar_surface_counts"]["earnings_calendar"] == 2
        assert surface["task_calendar_surface_counts"]["research_meeting"] == 1
        assert surface["task_calendar_surface_counts"]["risk_check"] == 3
        assert surface["events_with_time"] == 4
        assert surface["events_with_due_or_start"] == 4
        assert surface["events_with_reminders"] == 1
        assert surface["events_with_meeting_url"] == 1
        assert surface["events_with_project_or_calendar"] == 4
        assert surface["events_with_time_zone"] == 1
        assert surface["events_with_recurrence"] == 1
        assert surface["recurrence_frequency_counts"] == {"weekly": 1}
        assert surface["events_with_duration_minutes"] == 1
        assert surface["multi_day_event_count"] == 0
        assert surface["invalid_time_range_count"] == 0
        assert surface["events_with_checklist"] == 1
        assert surface["checklist_item_total"] == 3
        assert surface["checklist_item_completed_count"] == 2
        assert surface["checklist_item_pending_count"] == 1
        assert surface["average_checklist_completion_rate"] == 0.6667
        assert surface["tasks_with_incomplete_checklist"] == 1
        assert surface["collector_writes_wiki_directly"] is False
        proof = manifest["task_calendar_boundary_proof"]
        assert proof["proof_level"] == "authorized_task_calendar_with_time_quality"
        assert proof["event_count"] == 4
        assert proof["candidate_record_count"] == 5
        assert proof["matched_event_count"] == 4
        assert proof["filtered_candidate_count"] == 1
        assert proof["upstream_boundary"]["upstream_collector_counts"] == {"calendar": 1, "ticktick": 3}
        assert proof["upstream_boundary"]["source_platform_counts"] == {"feishu_calendar": 1, "ticktick": 3}
        assert proof["time_boundary"]["events_with_time"] == 4
        assert proof["time_boundary"]["events_with_due_or_start"] == 4
        assert proof["time_boundary"]["events_with_reminders"] == 1
        assert proof["time_boundary"]["events_with_meeting_url"] == 1
        assert proof["time_boundary"]["events_with_duration_minutes"] == 1
        assert proof["time_boundary"]["events_with_time_zone"] == 1
        assert proof["time_boundary"]["events_with_recurrence"] == 1
        assert proof["task_structure_boundary"]["events_with_checklist"] == 1
        assert proof["task_structure_boundary"]["checklist_item_total"] == 3
        assert proof["task_structure_boundary"]["checklist_item_completed_count"] == 2
        assert proof["task_structure_boundary"]["checklist_item_pending_count"] == 1
        assert proof["complete_task_list_claimed"] is False
        assert proof["complete_calendar_claimed"] is False
        assert proof["complete_task_calendar_context_claimed"] is False
        assert proof["direct_task_or_calendar_reconnect"] is False
        assert proof["requires_upstream_task_calendar_collector"] is True
        assert proof["task_calendar_boundary"]["task_calendar_surface_counts"]["risk_check"] == 3

        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["task-calendar-investor"]
        assert evidence_surface["task_calendar_surface_counts"]["trade_plan"] == 1
        assert evidence_surface["events_with_duration_minutes"] == 1
        assert evidence_surface["events_with_checklist"] == 1
        assert evidence_surface["generic_task_calendar_lens"] is True


def test_meeting_minutes_lens_keeps_investment_minutes_only() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "meeting-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "meeting:1",
                "collector": "meeting-artifacts",
                "source": "用户授权会议产物",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T11:00:00+08:00",
                "data": {
                    "artifact_type": "minutes",
                    "title": "半导体公司路演纪要",
                    "text_preview": "讨论财报、估值、风险点和买入框架。",
                    "participants": ["研究员A", "基金经理B"],
                },
                "raw_ref": {"path": "roadshow.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "meeting:2",
                "collector": "meeting-artifacts",
                "source": "用户授权会议产物",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T12:00:00+08:00",
                "collected_at": "2026-07-08T13:00:00+08:00",
                "data": {
                    "artifact_type": "minutes",
                    "title": "周五团建安排",
                    "text_preview": "讨论聚餐地点和出发时间。",
                },
                "raw_ref": {"path": "team.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "meeting-minutes",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T13:30:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "meeting-minutes" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 1
        assert lens_events[0]["data"]["payload"]["title"] == "半导体公司路演纪要"
        assert "matched_source_profile_terms" in lens_events[0]["data"]["classification"]["reasons"]


def test_meeting_minutes_lens_reports_meeting_surface_from_upstream_events() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "meeting-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "meeting-artifacts:roadshow",
                "collector": "meeting-artifacts",
                "source": "用户授权会议文件",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T11:00:00+08:00",
                "data": {
                    "artifact_type": "minutes",
                    "title": "半导体公司路演纪要",
                    "text_preview": "业绩路演讨论财报、估值、风险点和下一步跟进现金流假设。",
                    "participants": ["研究员A", "基金经理B"],
                    "meeting_url": "https://meeting.example/roadshow",
                    "attachments": [{"name": "roadshow-deck.pdf"}],
                    "action_items": ["下一步跟进现金流假设"],
                    "risk_items": ["估值风险点"],
                    "source_platform": "meeting-artifacts",
                },
                "raw_ref": {"path": "roadshow.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "feishu:expert-call",
                "collector": "feishu",
                "source": "飞书会议纪要",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T11:00:00+08:00",
                "data": {
                    "artifact_type": "transcript",
                    "title": "新能源专家电话会调研纪要",
                    "text_preview": "专家电话会调研供需拐点、价格弹性和下行风险。",
                    "participants": ["专家C"],
                    "meeting_url": "https://feishu.example/minutes/1",
                    "risk_items": ["价格弹性和下行风险"],
                    "source_platform": "feishu",
                },
                "raw_ref": {"path": "feishu-expert.json"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "dingtalk:ic",
                "collector": "dingtalk",
                "source": "钉钉协作导出",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T11:00:00+08:00",
                "collected_at": "2026-07-08T11:30:00+08:00",
                "data": {
                    "artifact_type": "meeting_ref",
                    "title": "投委会：组合调仓决策会",
                    "text_preview": "投资委员会记录决策点：加仓半导体、减仓高波动仓位，风控待办。",
                    "attendees": ["基金经理B", "风控D"],
                    "file_refs": [{"name": "ic-action-items.xlsx"}],
                    "decision_points": ["加仓半导体、减仓高波动仓位"],
                    "action_items": ["风控待办"],
                    "risk_items": ["高波动仓位"],
                    "source_platform": "dingtalk",
                },
                "raw_ref": {"path": "dingtalk-ic.json"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wecom:earnings",
                "collector": "wecom",
                "source": "企业微信会议导出",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T12:00:00+08:00",
                "collected_at": "2026-07-08T12:30:00+08:00",
                "data": {
                    "artifact_type": "recording_ref",
                    "title": "贵州茅台财报电话会 600519",
                    "text_preview": "业绩说明会复盘财报、现金流、毛利率和纪要结论。",
                    "participants": ["IR", "研究员A"],
                    "recording_refs": [{"url": "https://wecom.example/recording/1"}],
                    "mentioned_symbols": ["600519"],
                    "source_platform": "wecom",
                },
                "raw_ref": {"path": "wecom-earnings.json"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "meeting-artifacts:life",
                "collector": "meeting-artifacts",
                "source": "用户授权会议文件",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T13:00:00+08:00",
                "collected_at": "2026-07-08T14:00:00+08:00",
                "data": {
                    "artifact_type": "minutes",
                    "title": "周五团建安排",
                    "text_preview": "讨论聚餐地点和出发时间。",
                    "participants": ["同事E"],
                    "source_platform": "meeting-artifacts",
                },
                "raw_ref": {"path": "team.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["work_confidential"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "meeting-minutes",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T15:00:00+08:00",
        )

        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "meeting-minutes" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        upstream_ids = {event["raw_ref"]["upstream_event_id"] for event in lens_events}
        assert len(lens_events) == 4
        assert "meeting-artifacts:life" not in upstream_ids
        all_surfaces = {
            surface
            for event in lens_events
            for surface in event["data"]["classification"]["meeting_minutes_surfaces"]
        }
        assert {
            "roadshow_minutes",
            "research_meeting",
            "investment_committee",
            "expert_call",
            "earnings_call",
            "decision_point",
            "risk_discussion",
            "follow_up_action",
        }.issubset(all_surfaces)

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        surface = manifest["lens_surface_summary"]
        assert surface["event_count"] == 4
        assert surface["missing_expected_meeting_minutes_surfaces"] == []
        assert surface["meeting_minutes_surface_counts"]["roadshow_minutes"] == 1
        assert surface["meeting_minutes_surface_counts"]["investment_committee"] == 1
        assert surface["meeting_minutes_surface_counts"]["expert_call"] == 1
        assert surface["meeting_minutes_surface_counts"]["earnings_call"] == 1
        assert surface["upstream_collector_counts"] == {
            "dingtalk": 1,
            "feishu": 1,
            "meeting-artifacts": 1,
            "wecom": 1,
        }
        assert surface["kind_counts"] == {"note": 4}
        assert surface["source_platform_counts"] == {
            "dingtalk": 1,
            "feishu": 1,
            "meeting-artifacts": 1,
            "wecom": 1,
        }
        assert surface["participant_event_count"] == 4
        assert surface["participant_ref_count"] == 7
        assert surface["participant_role_event_count"] == 4
        assert surface["participant_role_counts"] == {
            "analyst": 2,
            "company_ir": 1,
            "expert": 1,
            "portfolio_manager": 2,
            "risk_control": 1,
        }
        assert surface["meeting_url_event_count"] == 2
        assert surface["attachment_ref_event_count"] == 2
        assert surface["recording_ref_event_count"] == 1
        assert surface["decision_point_event_count"] == 2
        assert surface["decision_point_count"] == 3
        assert surface["action_item_event_count"] == 2
        assert surface["action_item_count"] == 4
        assert surface["risk_item_event_count"] == 3
        assert surface["risk_item_count"] == 5
        assert surface["mentioned_symbol_event_count"] == 1
        assert surface["mentioned_symbol_count"] == 1
        assert surface["matched_symbol_event_count"] == 1
        assert surface["events_with_time"] == 4
        assert surface["collector_writes_wiki_directly"] is False
        proof = manifest["meeting_minutes_boundary_proof"]
        assert proof["proof_level"] == "authorized_meeting_minutes_with_decision_action_surface"
        assert proof["event_count"] == 4
        assert proof["candidate_record_count"] == 5
        assert proof["matched_event_count"] == 4
        assert proof["filtered_candidate_count"] == 1
        assert proof["upstream_boundary"]["upstream_collector_counts"] == {
            "dingtalk": 1,
            "feishu": 1,
            "meeting-artifacts": 1,
            "wecom": 1,
        }
        assert proof["meeting_context_boundary"]["participant_event_count"] == 4
        assert proof["meeting_context_boundary"]["participant_ref_count"] == 7
        assert proof["meeting_context_boundary"]["participant_role_event_count"] == 4
        assert proof["meeting_context_boundary"]["mentioned_symbol_event_count"] == 1
        assert proof["meeting_context_boundary"]["meeting_url_event_count"] == 2
        assert proof["meeting_context_boundary"]["attachment_ref_event_count"] == 2
        assert proof["meeting_context_boundary"]["recording_ref_event_count"] == 1
        assert proof["meeting_context_boundary"]["events_with_time"] == 4
        assert proof["decision_action_boundary"]["decision_point_event_count"] == 2
        assert proof["decision_action_boundary"]["action_item_event_count"] == 2
        assert proof["decision_action_boundary"]["risk_item_event_count"] == 3
        assert proof["complete_meeting_history_claimed"] is False
        assert proof["complete_workspace_claimed"] is False
        assert proof["complete_meeting_context_claimed"] is False
        assert proof["recording_body_collected_by_default"] is False
        assert proof["direct_meeting_platform_reconnect"] is False
        assert proof["requires_upstream_meeting_or_collaboration_collector"] is True
        assert proof["meeting_minutes_boundary"]["meeting_minutes_surface_counts"]["risk_discussion"] == 3

        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["meeting-minutes"]
        assert evidence_surface["meeting_minutes_surface_counts"]["risk_discussion"] == 3
        assert evidence_surface["action_item_event_count"] == 2
        assert evidence_surface["decision_point_event_count"] == 2
        assert evidence_surface["generic_meeting_lens"] is True


def test_wechat_article_favorites_lens_keeps_investment_articles_only() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "wechat-favorites-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:1",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T10:00:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "favorite",
                    "title": "半导体行业景气跟踪",
                    "source_account": "券商研究公众号",
                    "text_preview": "讨论财报、估值、风险点和安全边际。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/investment"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:2",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T11:00:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "favorite",
                    "title": "周末做饭清单",
                    "source_account": "生活号",
                    "text_preview": "采购食材和聚餐安排。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/life"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "wechat-article-favorites",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T11:30:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "wechat-article-favorites" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 1
        assert lens_events[0]["data"]["payload"]["title"] == "半导体行业景气跟踪"
        assert lens_events[0]["raw_ref"]["upstream_event_id"] == "wechat-favorites:1"


def test_wechat_article_favorites_lens_reports_article_surface_and_actions() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "wechat-favorites-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:research",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T10:00:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "favorite",
                    "title": "半导体行业景气跟踪",
                    "source_account": "券商研究公众号",
                    "source_account_type": "broker_research_account",
                    "url": "https://mp.weixin.qq.com/s/research",
                    "article_id": "research-001",
                    "tags": ["行业", "半导体"],
                    "symbols": ["688981"],
                    "favorite_reason": "跟踪半导体景气和估值假设",
                    "read_duration_seconds": 240,
                    "read_progress": "90%",
                    "engagement": {"read_count": 5600, "like_count": 120},
                    "text_preview": "证券研究深度报告讨论财报、现金流、估值、ROE、安全边际和风险提示。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/research"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:strategy",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T10:30:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "read",
                    "title": "市场策略周报：A股仓位与风格 000300",
                    "source_account": "财经策略号",
                    "source_account_type": "finance_media_account",
                    "url": "https://mp.weixin.qq.com/s/strategy",
                    "article_id": "strategy-001",
                    "tags": ["策略"],
                    "symbols": ["000300"],
                    "read_duration": "3分钟",
                    "read_progress": "60%",
                    "text_preview": "A股策略讨论仓位、风格、配置、宏观政策、利率和流动性。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/strategy"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:valuation",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T11:00:00+08:00",
                "collected_at": "2026-07-08T11:30:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "saved_file",
                    "title": "DCF估值模型笔记",
                    "source_account": "价值投资笔记",
                    "source_account_type": "investment_creator_account",
                    "url": "https://mp.weixin.qq.com/s/valuation",
                    "article_id": "valuation-001",
                    "tags": ["估值"],
                    "symbols": ["600519"],
                    "favorite_reason": "补充估值模型模板",
                    "text_preview": "估值模型拆解 DCF、PE、PB、ROE、目标价和安全边际。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/valuation"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:risk",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T12:00:00+08:00",
                "collected_at": "2026-07-08T12:30:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "share",
                    "title": "医药组合复盘：风险预警与调仓",
                    "source_account": "投资复盘号",
                    "source_account_type": "investment_creator_account",
                    "url": "https://mp.weixin.qq.com/s/risk",
                    "article_id": "risk-001",
                    "tags": ["组合", "风险"],
                    "symbols": ["300760"],
                    "share_target": "投研群",
                    "engagement": {"share_count": 32, "comment_count": 8},
                    "text_preview": "组合持仓复盘、回撤、止损、下行风险和调仓动作。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/risk"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "wechat-favorites:life",
                "collector": "wechat-favorites",
                "source": "微信收藏/公众号文章",
                "owner_scope": "personal",
                "kind": "file",
                "time": "2026-07-08T13:00:00+08:00",
                "collected_at": "2026-07-08T13:30:00+08:00",
                "data": {
                    "item_type": "public_account_article",
                    "action_type": "favorite",
                    "title": "周末做饭清单",
                    "source_account": "生活号",
                    "url": "https://mp.weixin.qq.com/s/life",
                    "text_preview": "采购食材和聚餐安排。",
                },
                "raw_ref": {"url": "https://mp.weixin.qq.com/s/life"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "wechat-article-favorites",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T14:00:00+08:00",
        )

        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "wechat-article-favorites" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        upstream_ids = {event["raw_ref"]["upstream_event_id"] for event in lens_events}
        assert len(lens_events) == 4
        assert "wechat-favorites:life" not in upstream_ids
        all_surfaces = {
            surface
            for event in lens_events
            for surface in event["data"]["classification"]["wechat_article_surfaces"]
        }
        assert {
            "broker_research_article",
            "company_fundamental_article",
            "market_strategy_article",
            "industry_theme_article",
            "valuation_method_article",
            "portfolio_case_article",
            "risk_warning_article",
            "macro_policy_article",
        }.issubset(all_surfaces)

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        surface = manifest["lens_surface_summary"]
        assert surface["event_count"] == 4
        assert surface["missing_expected_wechat_article_surfaces"] == []
        assert surface["action_type_counts"] == {
            "favorite": 1,
            "read": 1,
            "saved_file": 1,
            "share": 1,
        }
        assert surface["upstream_collector_counts"] == {"wechat-favorites": 4}
        assert surface["source_account_type_counts"]["broker_research_account"] == 1
        assert surface["source_account_type_counts"]["finance_media_account"] == 1
        assert surface["source_account_type_counts"]["investment_creator_account"] == 2
        assert surface["source_account_count"] == 4
        assert surface["public_account_article_count"] == 4
        assert surface["matched_symbol_event_count"] == 4
        assert surface["symbol_event_count"] == 4
        assert surface["symbol_count"] == 4
        assert surface["events_with_url"] == 4
        assert surface["events_with_article_id"] == 4
        assert surface["events_with_tags"] == 4
        assert surface["events_with_text"] == 4
        assert surface["events_with_action_time"] == 4
        assert surface["events_with_favorite_reason"] == 2
        assert surface["events_with_share_target"] == 1
        assert surface["events_with_read_duration"] == 2
        assert surface["events_with_read_progress"] == 2
        assert surface["events_with_engagement"] == 2
        assert surface["average_read_duration_seconds"] == 210
        assert surface["average_read_progress"] == 0.75
        assert surface["collector_writes_wiki_directly"] is False
        proof = manifest["wechat_article_boundary_proof"]
        assert proof["proof_level"] == "authorized_wechat_articles_with_behavior_surface"
        assert proof["event_count"] == 4
        assert proof["candidate_record_count"] == 5
        assert proof["matched_event_count"] == 4
        assert proof["filtered_candidate_count"] == 1
        assert proof["upstream_boundary"]["upstream_collector_counts"] == {"wechat-favorites": 4}
        assert proof["upstream_boundary"]["item_type_counts"] == {"public_account_article": 4}
        assert proof["article_action_boundary"]["action_type_counts"] == {
            "favorite": 1,
            "read": 1,
            "saved_file": 1,
            "share": 1,
        }
        assert proof["article_action_boundary"]["source_account_count"] == 4
        assert proof["article_action_boundary"]["public_account_article_count"] == 4
        assert proof["article_action_boundary"]["matched_symbol_event_count"] == 4
        assert proof["article_action_boundary"]["symbol_event_count"] == 4
        assert proof["article_action_boundary"]["symbol_count"] == 4
        assert proof["content_pointer_boundary"]["events_with_url"] == 4
        assert proof["content_pointer_boundary"]["events_with_article_id"] == 4
        assert proof["content_pointer_boundary"]["events_with_source_account"] == 4
        assert proof["content_pointer_boundary"]["events_with_tags"] == 4
        assert proof["content_pointer_boundary"]["events_with_text"] == 4
        assert proof["content_pointer_boundary"]["events_with_action_time"] == 4
        assert proof["behavior_boundary"]["events_with_favorite_reason"] == 2
        assert proof["behavior_boundary"]["events_with_share_target"] == 1
        assert proof["behavior_boundary"]["events_with_read_duration"] == 2
        assert proof["behavior_boundary"]["events_with_read_progress"] == 2
        assert proof["behavior_boundary"]["events_with_engagement"] == 2
        assert proof["behavior_boundary"]["average_read_duration_seconds"] == 210
        assert proof["behavior_boundary"]["average_read_progress"] == 0.75
        assert proof["complete_wechat_favorites_claimed"] is False
        assert proof["complete_wechat_read_history_claimed"] is False
        assert proof["public_account_full_crawl_claimed"] is False
        assert proof["public_article_body_mirrored"] is False
        assert proof["direct_wechat_reconnect"] is False
        assert proof["requires_upstream_wechat_favorites_collector"] is True
        assert proof["wechat_article_boundary"]["wechat_article_surface_counts"]["risk_warning_article"] == 2

        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["wechat-article-favorites"]
        assert evidence_surface["wechat_article_surface_counts"]["risk_warning_article"] == 2
        assert evidence_surface["generic_wechat_article_lens"] is True


def test_social_investment_influence_lens_keeps_investment_activity_only() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "social-events.jsonl"
        out_dir = root / "out"
        events = [
            {
                "schema": "collectorx.event.v1",
                "id": "social-activity:1",
                "collector": "social-activity",
                "source": "社交平台用户授权活动",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T11:00:00+08:00",
                "data": {
                    "platform": "bilibili",
                    "action_type": "watch",
                    "title": "半导体投资复盘：财报、估值和安全边际",
                    "creator": "财经博主A",
                    "tags": ["股票", "半导体"],
                    "social_topics": ["industry_theme", "company_fundamental", "risk_control"],
                    "primary_social_topic": "industry_theme",
                    "content_preview": "实盘复盘行业景气、估值和风险点。",
                },
                "raw_ref": {"path": "bilibili.csv", "row": 1},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "social-activity:2",
                "collector": "social-activity",
                "source": "社交平台用户授权活动",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T11:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "platform": "bilibili",
                    "action_type": "like",
                    "title": "游戏直播剪辑",
                    "creator": "娱乐UP主",
                    "tags": ["游戏"],
                    "content_preview": "娱乐内容。",
                },
                "raw_ref": {"path": "bilibili.csv", "row": 2},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_message", "contact"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "social-investment-influence",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T12:30:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "social-investment-influence" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 1
        assert lens_events[0]["data"]["payload"]["title"].startswith("半导体投资复盘")
        assert lens_events[0]["raw_ref"]["upstream_event_id"] == "social-activity:1"
        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        surface = manifest["lens_surface_summary"]
        assert surface["social_topic_counts"]["industry_theme"] == 1
        assert surface["social_topic_counts"]["company_fundamental"] == 1
        assert surface["social_topic_counts"]["risk_control"] == 1
        assert surface["platform_counts"] == {"bilibili": 1}
        assert surface["action_counts"] == {"watch": 1}
        assert surface["usable_as_investment_conclusion"] is False
        proof = manifest["social_influence_boundary_proof"]
        assert proof["proof_level"] == "medium_partial_social_influence_boundary"
        assert proof["weak_evidence_only"] is True
        assert proof["requires_corroboration"] is True
        assert proof["can_claim_investment_conclusion"] is False
        assert proof["social_topic_boundary"]["observed_social_topics"] == [
            "industry_theme",
            "company_fundamental",
            "risk_control",
        ]
        assert proof["platform_action_boundary"]["platform_counts"] == {"bilibili": 1}
        assert proof["creator_content_boundary"]["creator_event_count"] == 1
        assert proof["false_claims"]["investment_conclusion_claimed"] is False
        assert proof["false_claims"]["full_content_mirrored"] is False
        assert "strong_trade_research_corroboration_missing" in proof["completion_blockers"]
        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["social-investment-influence"]
        assert evidence_surface["requires_corroboration"] is True
        assert evidence_surface["social_topic_counts"]["industry_theme"] == 1
        evidence_proof = evidence["coverage_summary"]["source_boundary_proof_summary"]["social-investment-influence"]
        assert evidence_proof["weak_evidence_only"] is True
        assert evidence_proof["can_claim_investment_conclusion"] is False


def test_investment_notes_lens_reports_note_type_surface_from_notes_events() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source_path = root / "notes-events.jsonl"
        out_dir = root / "out"
        notes_events = [
            {
                "schema": "collectorx.event.v1",
                "id": "notes:review",
                "collector": "notes",
                "source": "Obsidian vault",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T09:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "obsidian",
                    "title": "贵州茅台交易复盘",
                    "path": "reviews/600519.md",
                    "content_preview": "复盘 600519 买入理由、错误归因、仓位和交易纪律。",
                    "content_length": 28,
                    "content_included": False,
                    "tags": ["复盘", "投资"],
                },
                "raw_ref": {"source_app": "obsidian", "path": "reviews/600519.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_note"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "notes:rules",
                "collector": "notes",
                "source": "Notion API",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T10:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "notion",
                    "title": "交易规则库",
                    "path": "notion-page-rules",
                    "content_preview": "规则库 checklist 买入框架、卖出框架、仓位控制、止损和止盈。",
                    "content_length": 35,
                    "content_included": False,
                    "tags": ["规则库"],
                },
                "raw_ref": {"source_app": "notion", "path": "notion-page-rules"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_note"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "notes:valuation",
                "collector": "notes",
                "source": "有道云笔记导出",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T11:00:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "youdao",
                    "title": "半导体估值假设",
                    "path": "youdao/semiconductor.md",
                    "content_preview": "估值假设 DCF PE PB ROE 现金流 安全边际 财报 跟踪。",
                    "content_length": 40,
                    "content_included": False,
                    "tags": ["估值"],
                },
                "raw_ref": {"source_app": "youdao", "path": "youdao/semiconductor.md"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_note"]},
            },
            {
                "schema": "collectorx.event.v1",
                "id": "notes:life",
                "collector": "notes",
                "source": "Evernote export",
                "owner_scope": "personal",
                "kind": "note",
                "time": "2026-07-08T11:30:00+08:00",
                "collected_at": "2026-07-08T12:00:00+08:00",
                "data": {
                    "source_app": "evernote",
                    "title": "周末安排",
                    "path": "life/weekend.enex",
                    "content_preview": "周末买菜、整理房间、约朋友吃饭。",
                    "content_length": 18,
                    "content_included": False,
                },
                "raw_ref": {"source_app": "evernote", "path": "life/weekend.enex"},
                "privacy": {"sensitive": True, "local_only": True, "contains": ["personal_note"]},
            },
        ]
        source_path.write_text("\n".join(json.dumps(event, ensure_ascii=False) for event in notes_events) + "\n", encoding="utf-8")
        run_cli(
            "collect",
            "--source",
            "investment-notes",
            "--input",
            str(source_path),
            "--out-dir",
            str(out_dir),
            "--collected-at",
            "2026-07-08T12:30:00+08:00",
        )
        lens_events = [
            json.loads(line)
            for line in (out_dir / "lake" / "investment-notes" / "events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        assert len(lens_events) == 3
        assert {event["raw_ref"]["upstream_event_id"] for event in lens_events} == {"notes:review", "notes:rules", "notes:valuation"}
        classifications = [event["data"]["classification"] for event in lens_events]
        assert any("review_note" in item["investment_note_types"] for item in classifications)
        assert any("rules_library" in item["investment_note_types"] for item in classifications)
        assert any("trade_checklist" in item["investment_note_types"] for item in classifications)
        assert any("valuation_assumption" in item["investment_note_types"] for item in classifications)
        assert any("research_note" in item["investment_note_types"] for item in classifications)

        manifest = json.loads((out_dir / "manifest.json").read_text(encoding="utf-8"))
        surface = manifest["lens_surface_summary"]
        assert surface["event_count"] == 3
        assert surface["source_app_counts"] == {"obsidian": 1, "notion": 1, "youdao": 1}
        assert surface["upstream_collector_counts"] == {"notes": 3}
        assert surface["missing_expected_investment_note_types"] == []
        assert surface["investment_note_type_counts"]["review_note"] == 1
        assert surface["investment_note_type_counts"]["rules_library"] == 2
        assert surface["investment_note_type_counts"]["trade_checklist"] == 2
        assert surface["investment_note_type_counts"]["valuation_assumption"] == 1
        assert surface["investment_note_type_counts"]["research_note"] == 1
        assert surface["full_content_event_count"] == 0
        assert surface["preview_only_event_count"] == 3
        assert surface["tagged_event_count"] == 3
        assert surface["collector_writes_wiki_directly"] is False
        proof = manifest["investment_note_boundary_proof"]
        assert proof["proof_level"] == "authorized_investment_notes_preview_only"
        assert proof["event_count"] == 3
        assert proof["candidate_record_count"] == 4
        assert proof["matched_event_count"] == 3
        assert proof["filtered_candidate_count"] == 1
        assert proof["complete_notes_vault_claimed"] is False
        assert proof["complete_note_context_claimed"] is False
        assert proof["direct_notes_reconnect"] is False
        assert proof["requires_upstream_notes_collector"] is True
        assert proof["content_boundary"]["full_content_event_count"] == 0
        assert proof["content_boundary"]["preview_only_event_count"] == 3
        assert proof["content_boundary"]["tagged_event_count"] == 3
        assert proof["content_boundary"]["path_event_count"] == 3
        assert proof["note_boundary"]["source_app_counts"] == {"obsidian": 1, "notion": 1, "youdao": 1}
        assert proof["note_boundary"]["investment_note_type_counts"]["rules_library"] == 2

        evidence = json.loads((out_dir / "investor_wiki_evidence.v1.json").read_text(encoding="utf-8"))
        evidence_surface = evidence["coverage_summary"]["source_surface_summary"]["investment-notes"]
        assert evidence_surface["investment_note_type_counts"]["valuation_assumption"] == 1
        assert evidence_surface["generic_notes_lens"] is True


if __name__ == "__main__":
    test_list_sources_contains_all_priorities()
    test_collect_xueqiu_csv_outputs_event_and_evidence()
    test_email_research_scope_policy_does_not_filter_other_lenses()
    test_collect_without_input_writes_gap_event()
    test_wechat_lens_keeps_only_investment_dialogue()
    test_wechat_lens_source_policy_allows_and_denies_chats_and_senders()
    test_wechat_lens_source_policy_filtered_all_has_explicit_gap()
    test_lens_without_investment_match_does_not_fill_wiki_coverage()
    test_email_research_reads_upstream_collectorx_event()
    test_email_research_matches_research_attachment_filename()
    test_email_research_reports_surface_and_boundary_proof()
    test_email_research_scope_policy_filters_authorized_research_mail()
    test_email_research_scope_policy_filtered_all_gap()
    test_research_documents_extracts_office_and_pdf_content_when_authorized()
    test_research_documents_extracts_legacy_xls_and_pptx_when_authorized()
    test_research_documents_binary_xls_without_xlrd_records_explicit_failure()
    test_research_documents_without_include_content_keeps_binary_metadata_only()
    test_research_documents_filters_broad_titles_and_skips_unsupported_files()
    test_research_documents_scope_policy_filters_authorized_surface()
    test_research_documents_scope_policy_filtered_all_has_explicit_gap()
    test_research_documents_image_ocr_requires_explicit_adapter_authorization()
    test_research_documents_image_ocr_extracts_when_explicitly_authorized()
    test_research_documents_limit_records_path_truncation()
    test_task_calendar_lens_keeps_investment_task_and_calendar_only()
    test_task_calendar_lens_reports_planning_surface_from_upstream_events()
    test_meeting_minutes_lens_keeps_investment_minutes_only()
    test_meeting_minutes_lens_reports_meeting_surface_from_upstream_events()
    test_wechat_article_favorites_lens_keeps_investment_articles_only()
    test_wechat_article_favorites_lens_reports_article_surface_and_actions()
    test_social_investment_influence_lens_keeps_investment_activity_only()
    test_investment_notes_lens_reports_note_type_surface_from_notes_events()
    print("investor-source-collectors tests passed.")
